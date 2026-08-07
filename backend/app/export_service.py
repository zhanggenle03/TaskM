"""任务导出服务 —— 生成 DOCX 文档 + 附件 ZIP 包"""

import os
import io
import re
import zipfile
from datetime import datetime
from typing import Optional, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from sqlalchemy.orm import Session, joinedload
from .database import (
    Communication, Attachment, Contact, StatusPool,
    TagPool, TaskTag, Project, Requirement,
    UPLOAD_DIR, derive_task_status
)

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg', '.ico'}

TASK_ATTR_OPTIONS = {
    'title': '任务名称', 'display_id': '显示ID', 'status': '状态',
    'priority': '优先级', 'due_date': '截止日期', 'description': '描述',
    'created_at': '创建时间', 'updated_at': '更新时间',
    'contacts': '对接人', 'tags': '标签',
    'linked_requirements': '关联需求',
}

PRIORITY_LABELS = {'low': '低', 'normal': '普通', 'high': '高', 'urgent': '紧急'}

# 公文风格字号常量
FONT_FAMILY = '仿宋'
FONT_FAMILY_HEADING = '黑体'
# 西文（英文/数字）字体：标题与内容统一使用新罗马，中文仍用上述中文字体
LATIN_FONT = 'Times New Roman'
BODY_SIZE = Pt(12)        # 小四
SMALL_SIZE = Pt(10)       # 五号
HEADING1_SIZE = Pt(16)    # 三号
HEADING2_SIZE = Pt(14)    # 四号
HEADING3_SIZE = Pt(13)    # 小三号
HEADING4_SIZE = Pt(12)    # 小四号
HEADING5_SIZE = Pt(11)    # 五号
HEADING6_SIZE = Pt(10.5)  # 小五号
TITLE_SIZE = Pt(22)       # 二号
SUBTITLE_SIZE = Pt(14)    # 四号


def _set_cell_shading(cell, color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, size=BODY_SIZE, bold=False, color=None, font_name=FONT_FAMILY):
    """为 run 设置标准的公文风格字体：西文（英文/数字）统一新罗马，中文用 font_name"""
    run.bold = bold
    run.font.size = size
    # 西文字体固定为新罗马（w:ascii + w:hAnsi），英文/数字标题与内容统一
    run.font.name = LATIN_FONT
    if color:
        run.font.color.rgb = color
    # 设置东亚字体
    r_elem = run._element
    rPr = r_elem.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    return run


def _add_run(paragraph, text, size=BODY_SIZE, bold=False, color=None, font_name=FONT_FAMILY):
    """添加 run 并设置公文风格字体"""
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color, font_name=font_name)
    return run


def _new_paragraph(doc, text='', size=BODY_SIZE, bold=False, color=None,
                   font_name=FONT_FAMILY, alignment=None,
                   before=0, after=0, first_line_indent=None):
    """创建新段落并设置公文格式"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if text:
        _add_run(p, text, size=size, bold=bold, color=color, font_name=font_name)

    # 设置段落间距
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    if before:
        spacing.set(qn('w:before'), str(before))
    if after:
        spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), '360')  # 1.5 倍行距
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

    # 首行缩进
    if first_line_indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), str(first_line_indent))
        pPr.append(ind)

    return p


def _add_hyperlink(paragraph, text, url, size=Pt(12)):
    """在段落中添加超链接（蓝色+下划线）"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    r_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), LATIN_FONT)
    rFonts.set(qn('w:hAnsi'), LATIN_FONT)
    rFonts.set(qn('w:eastAsia'), FONT_FAMILY)
    rPr.append(rFonts)
    # 字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size.pt * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size.pt * 2))
    rPr.append(szCs)
    # 蓝色
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0066CC')
    rPr.append(c)
    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    r_elem.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    r_elem.append(t_elem)
    hyperlink.append(r_elem)
    paragraph._p.append(hyperlink)
    return hyperlink


def _set_heading_style(doc, level, font_name=FONT_FAMILY_HEADING, size=HEADING1_SIZE, line=360):
    """设置标题样式（line 为行距：360=1.5 倍，240=单倍，单位 1/240 行）。
    西文（英文/数字）统一新罗马，中文用 font_name（黑体）。"""
    style = doc.styles[f'Heading {level}']
    style.font.name = LATIN_FONT
    style.font.size = size
    style.font.bold = True
    # 显式清除斜体：python-docx 默认模板中 Heading 4/5/6 自带斜体，
    # 若不覆盖，任务导出里四级标题（内容 h2 -> Heading4）会渲染成斜体
    style.font.italic = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 段落格式
    pPr = style.element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), '200')
    spacing.set(qn('w:after'), '200')
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')


def _ml(ilvl, fmt, text, start=1, suff='tab', lvl_font=None, lvl_size_pt=None):
    """构造一个多级编号的 w:lvl 元素。

    按 OOXML CT_Lvl 顺序：start -> numFmt -> suff -> lvlText -> lvlJc -> rPr。
    suff 决定编号后的分隔符：tab(默认，Word 会插入制表符) / space(空格) / nothing(无)。
    lvl_font/lvl_size_pt：给序号文本显式设置样式——西文(英文/数字)新罗马 + 中文 lvl_font
    + 加粗 + 字号，使序号与标题文字保持一致（Word 默认序号用普通字体，不与标题同步）。
    注意：lvlText 中 %1 恒指第一级计数、%2 指第二级、%3 指第三级，依此类推；
          不要用 %1 表示"本级"，否则所有同级标题会显示同一个上级编号。
    """
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), str(ilvl))
    for tag, val in [('start', str(start)), ('numFmt', fmt)]:
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:val'), val)
        lvl.append(el)
    suff_el = OxmlElement('w:suff')
    suff_el.set(qn('w:val'), suff)
    lvl.append(suff_el)
    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), text)
    lvl.append(lvlText)
    jc = OxmlElement('w:lvlJc')
    jc.set(qn('w:val'), 'left')
    lvl.append(jc)
    # 序号文本样式（rPr 位于 lvlJc 之后）
    if lvl_font and lvl_size_pt:
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), LATIN_FONT)
        rFonts.set(qn('w:hAnsi'), LATIN_FONT)
        rFonts.set(qn('w:eastAsia'), lvl_font)
        rPr.append(rFonts)
        rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(lvl_size_pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(lvl_size_pt * 2)))
        rPr.append(szCs)
        lvl.append(rPr)
    return lvl


def _setup_numbering(doc):
    """
    建立任务导出的【结构标题】编号（公文中文风，2 级）：
      一级（章节）：一、 二、 三、   -> 任务基本信息 / 沟通记录
      二级（每条记录）： （一） （二） （三） -> 各条沟通记录
    返回 numId=99，供 _apply_numbering 使用。
    """
    numbering_part = doc.part.numbering_part
    numbering = numbering_part.element
    for n in numbering.findall(qn('w:num')):
        if n.get(qn('w:numId')) == '99':
            return 99
    ab = OxmlElement('w:abstractNum')
    ab.set(qn('w:abstractNumId'), '99')
    # 一级：一、后面直接接标题（suff=nothing，顿号即分隔，无制表符）
    ab.append(_ml(0, 'chineseCounting', '%1、', suff='nothing',
                  lvl_font=FONT_FAMILY_HEADING, lvl_size_pt=HEADING1_SIZE.pt))
    # 二级：%2=本级计数（每个一级标题后从（一）重新开始），编号后接空格（非制表符）
    ab.append(_ml(1, 'chineseCounting', '（%2）', suff='space',
                  lvl_font=FONT_FAMILY_HEADING, lvl_size_pt=HEADING2_SIZE.pt))
    numbering.append(ab)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '99')
    ref = OxmlElement('w:abstractNumId')
    ref.set(qn('w:val'), '99')
    num.append(ref)
    numbering.append(num)
    return 99


def _setup_content_numbering(doc, num_id):
    """
    建立任务沟通【内容标题】的独立编号（阿拉伯风，3 级），与结构标题的中文编号分离：
      一级（内容 h1 -> Heading3）：1
      二级（内容 h2 -> Heading4）：1.1
      三级（内容 h3 -> Heading5）：(1)
    每条沟通记录使用独立实例（num_id 不同），编号从 1 重新开始。
    返回 num_id。
    """
    numbering_part = doc.part.numbering_part
    numbering = numbering_part.element
    for n in numbering.findall(qn('w:num')):
        if n.get(qn('w:numId')) == str(num_id):
            return num_id
    ab_id = str(1000 + num_id)
    ab = OxmlElement('w:abstractNum')
    ab.set(qn('w:abstractNumId'), ab_id)
    # 内容标题编号后统一接空格（suff=space，非制表符）；%3 为本级(h3)计数 -> (1)(2)...
    # 序号样式与内容标题一致：黑体 + 小四(12pt) + 加粗 + 西文新罗马
    ab.append(_ml(0, 'decimal', '%1', suff='space',
                  lvl_font=FONT_FAMILY_HEADING, lvl_size_pt=BODY_SIZE.pt))      # 1
    ab.append(_ml(1, 'decimal', '%1.%2', suff='space',
                  lvl_font=FONT_FAMILY_HEADING, lvl_size_pt=BODY_SIZE.pt))      # 1.1
    ab.append(_ml(2, 'decimal', '(%3)', suff='space',
                  lvl_font=FONT_FAMILY_HEADING, lvl_size_pt=BODY_SIZE.pt))      # (1)
    numbering.append(ab)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    ref = OxmlElement('w:abstractNumId')
    ref.set(qn('w:val'), ab_id)
    num.append(ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id, ilvl):
    """给已添加的段落应用自动编号"""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def _add_h1(doc, text, num_id, ilvl=0):
    """添加一级标题并自动分页（前置分页符 + Heading 1 + 自动编号）"""
    doc.add_page_break()
    h = doc.add_heading(text, level=1)
    _apply_numbering(h, num_id, ilvl)
    return h


def build_export_data(
    db: Session,
    project_id: str,
    task_id: str,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    fields: Optional[List[str]] = None,
    comm_ids: Optional[List[int]] = None,
) -> dict:
    """查询并整理导出数据"""
    from .database import resolve_project, resolve_task

    proj = resolve_project(db, project_id)
    task = resolve_task(db, proj.id, task_id)

    derived_status = derive_task_status(db, task.id)
    if derived_status is not None:
        task.status_id = derived_status

    project_info = db.query(Project).filter(Project.id == proj.id).first()
    status_pool = {s.id: s for s in db.query(StatusPool).filter(StatusPool.project_id == proj.id).all()}

    # 查询关联需求
    linked_requirements = []
    for req in (task.linked_requirements or []):
        linked_requirements.append({
            'id': req.id,
            'title': req.title,
            'display_id': req.display_id,
            'status': req.status,
            'priority': req.priority,
        })

    tag_rows = db.query(TaskTag).filter(TaskTag.task_id == task.id).all()
    tag_ids = [tr.tag_id for tr in tag_rows]
    tag_info = {}
    if tag_ids:
        for t in db.query(TagPool).filter(TagPool.id.in_(tag_ids)).all():
            tag_info[t.id] = t

    comm_query = db.query(Communication).filter(
        Communication.task_id == task.id
    ).order_by(Communication.comm_at, Communication.id)
    communications = comm_query.all()

    if start_date:
        try:
            sd = datetime.strptime(start_date, '%Y-%m-%d')
            communications = [c for c in communications if c.comm_at >= sd]
        except ValueError:
            pass
    if end_date:
        try:
            ed = datetime.strptime(end_date, '%Y-%m-%d').replace(hour=23, minute=59, second=59)
            communications = [c for c in communications if c.comm_at <= ed]
        except ValueError:
            pass
    if comm_ids:
        comm_id_set = set(comm_ids)
        communications = [c for c in communications if c.id in comm_id_set]

    contact_ids = set()
    for c in communications:
        for cc in c.communication_contacts:
            contact_ids.add(cc.contact_id)
    contact_map = {}
    if contact_ids:
        for ct in db.query(Contact).filter(Contact.id.in_(contact_ids)).all():
            contact_map[ct.id] = ct

    task_contacts = task.contacts or []
    tag_names = [tag_info[tid].name for tid in tag_ids if tid in tag_info]

    comm_list = []
    for c in communications:
        # 跳过"创建任务"系统初始化记录（无实际沟通内容），保持导出干净
        if c.old_status_id is None and (c.content or '').startswith('创建任务，初始状态：'):
            continue
        comm_contacts = []
        for cc in c.communication_contacts:
            ct = contact_map.get(cc.contact_id)
            if ct:
                comm_contacts.append(ct.name)

        att_list = []
        for att in (c.attachments or []):
            full_path = os.path.join(UPLOAD_DIR, proj.display_id, task.display_id, f'comm_{c.id}', att.filename)
            ext = os.path.splitext(att.original_filename)[1].lower()
            att_list.append({
                'id': att.id,
                'original_filename': att.original_filename,
                'file_size': att.file_size,
                'is_image': ext in IMAGE_EXTENSIONS,
                'full_path': full_path if os.path.isfile(full_path) else None,
            })

        old_status_name = status_pool[c.old_status_id].name if c.old_status_id and c.old_status_id in status_pool else ''
        new_status_name = status_pool[c.new_status_id].name if c.new_status_id and c.new_status_id in status_pool else ''

        comm_list.append({
            'id': c.id,
            'comm_at': c.comm_at,
            'comm_type': c.comm_type,
            'content': c.content,
            'contacts': comm_contacts,
            'old_status_id': c.old_status_id,
            'new_status_id': c.new_status_id,
            'subject': c.subject or '',
            'old_status_name': old_status_name,
            'new_status_name': new_status_name,
            'attachments': att_list,
        })

    task_attrs = {
        'title': task.title,
        'display_id': task.display_id,
        'status': status_pool[task.status_id].name if task.status_id and task.status_id in status_pool else '',
        'priority': PRIORITY_LABELS.get(task.priority, task.priority),
        'due_date': task.due_date.strftime('%Y-%m-%d') if task.due_date else '未设置',
        'description': task.description or '暂无描述',
        'created_at': task.created_at.strftime('%Y-%m-%d %H:%M') if task.created_at else '',
        'updated_at': task.updated_at.strftime('%Y-%m-%d %H:%M') if task.updated_at else '',
        'contacts': ', '.join([ct.name for ct in task_contacts]) if task_contacts else '无',
        'tags': ', '.join(tag_names) if tag_names else '无',
    }

    return {
        'project_name': project_info.name if project_info else '',
        'project_display_id': proj.display_id,
        'proj': proj,
        'task_attrs': task_attrs,
        'linked_requirements': linked_requirements,
        'communications': comm_list,
    }


def _sanitize_filename(name: str) -> str:
    """清理文件名中的非法字符（Windows 不允许的字符 -> 下划线）"""
    illegal = r'<>:"/\\|?*'
    for ch in illegal:
        name = name.replace(ch, '_')
    return name.strip()


def _add_task_info_table(doc: Document, task_attrs: dict, selected_fields: List[str], num_id: int = 99):
    """添加任务基本信息（表格）"""
    h = _add_h1(doc, '任务基本信息', num_id, 0)

    field_order = ['title', 'display_id', 'status', 'priority', 'due_date',
                   'description', 'contacts', 'tags', 'created_at', 'updated_at']
    field_labels = {
        'title': '任务名称', 'display_id': '显示ID', 'status': '状态',
        'priority': '优先级', 'due_date': '截止日期', 'description': '描述',
        'contacts': '对接人', 'tags': '标签', 'created_at': '创建时间', 'updated_at': '更新时间'
    }

    items = [(field_labels[k], task_attrs[k]) for k in field_order
             if k in selected_fields and k in task_attrs]

    if not items:
        return

    _apply_table_widths(doc, items)

    doc.add_paragraph()


def _apply_table_widths(doc, items, col1_twips=2448):
    """创建两列信息表，使用 Table Grid 样式 + 精确列宽控制。
    被任务导出和需求导出共用，确保两者表格样式 100% 一致。"""
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    twips = int(usable * 1440 / 914400)
    col2_twips = twips - col1_twips

    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 填充内容
    for i, (label, value) in enumerate(items):
        cell_label = table.rows[i].cells[0]
        _set_cell_shading(cell_label, 'F2F2F2')
        _add_run(cell_label.paragraphs[0], label, size=BODY_SIZE, bold=True,
                 font_name=FONT_FAMILY_HEADING)
        cell_value = table.rows[i].cells[1]
        _add_run(cell_value.paragraphs[0], value, size=BODY_SIZE)

    _fix_table_layout(table, twips, col1_twips, col2_twips)

    return table


def _fix_table_layout(table, twips, col1_twips, col2_twips):
    """修正表格的宽度属性：tblW、tblGrid、tcW，并锁定 fixed 布局。
    在 Table Grid 样式基础上精确覆盖列宽，不破坏边框等样式继承。"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))

    # 1. 修正 tblW：移除所有旧值，插入一个固定宽度（位置必须在 tblStyle 之后）
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(twips))
    tblW.set(qn('w:type'), 'dxa')
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblStyle.addnext(tblW)
    else:
        tblPr.insert(0, tblW)

    # 2. 添加 tblLayout type=fixed（确保 Word 不自动调整列宽）
    old_layout = tblPr.find(qn('w:tblLayout'))
    if old_layout is not None:
        tblPr.remove(old_layout)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    # tblLayout 应在 tblBorders 之后
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblBorders.addnext(tblLayout)
    else:
        tblPr.append(tblLayout)

    # 3. 重建 tblGrid
    old_grid = tbl.find(qn('w:tblGrid'))
    if old_grid is not None:
        tbl.remove(old_grid)
    new_grid = OxmlElement('w:tblGrid')
    for gw in [col1_twips, col2_twips]:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(gw))
        new_grid.append(gc)
    first_tr = tbl.find(qn('w:tr'))
    if first_tr is not None:
        tbl.insert(list(tbl).index(first_tr), new_grid)

    # 4. 逐单元格锁定 tcW
    target_widths = [col1_twips, col2_twips]
    for row in tbl.findall(qn('w:tr')):
        tcs = row.findall(qn('w:tc'))
        for ci, tc in enumerate(tcs):
            if ci >= len(target_widths):
                break
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            old_tcW = tcPr.find(qn('w:tcW'))
            if old_tcW is not None:
                tcPr.remove(old_tcW)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(target_widths[ci]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.insert(0, tcW)


def _format_size(bytes_val: int) -> str:
    if bytes_val > 1024 * 1024:
        return f'{bytes_val / 1024 / 1024:.1f}MB'
    elif bytes_val > 1024:
        return f'{bytes_val / 1024:.1f}KB'
    return f'{bytes_val}B'


# 自动生成的状态变更文本前缀（与 routers/tasks.py 中 update_task 生成格式一致）
_AUTO_STATUS_PREFIXES = ('状态变更：', '状态变更为：', '状态：')


def _is_auto_status_text(content: str, old_status_id, new_status_id) -> bool:
    """判断沟通内容是否为系统自动生成的状态变更文本（与状态行重复，导出/展示时应隐藏）。"""
    if not (old_status_id or new_status_id):
        return False
    if not content:
        return False
    c = content.strip()
    return any(c.startswith(p) for p in _AUTO_STATUS_PREFIXES)


_HTML_TAG_RE = re.compile(r'<([a-zA-Z][a-zA-Z0-9]*)\b[^>]*>', re.I)


def _looks_like_html(text: str) -> bool:
    """粗略判断内容是否为 HTML（含标签），用于决定是否走富文本渲染。"""
    return bool(_HTML_TAG_RE.search(text or ''))


def generate_export_package(
    db: Session,
    project_id: str,
    task_id: str,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    fields: Optional[List[str]] = None,
    comm_ids: Optional[List[int]] = None,
    comm_minimal: bool = False,
) -> bytes:
    """
    生成导出包（ZIP 文件），内含：
    1. DOCX 说明文档（含内嵌图片）
    2. attachments/{记录序号}/ 目录（每个沟通记录的附件各自放在对应编号的文件夹）
    """
    all_field_keys = list(TASK_ATTR_OPTIONS.keys())
    selected_fields = fields if fields is not None else all_field_keys

    data = build_export_data(db, project_id, task_id, start_date, end_date, selected_fields, comm_ids)
    task_attrs = data['task_attrs']
    communications = data['communications']
    project_name = data['project_name']
    project_display_id = data['project_display_id']
    proj = data['proj']

    # ======================== 生成 DOCX ========================
    doc = Document()

    # 建立自动编号定义（doc.add_heading 之前调用）
    num_id = _setup_numbering(doc)

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = LATIN_FONT
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY)
    # Normal 段落间距
    pPr = style.element.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:line'), '360')
    pSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpacing)

    # 设置标题样式（1-6 级均设为黑体、黑色、加粗，避免 Word 默认模板中
    # Heading 3+ 的蓝色/青色；与需求导出（routers/requirements.py）保持一致）
    # 任务文档标题统一单倍行距（line=240）；需求导出仍用默认 360（1.5 倍）
    for _lv, _sz in {
        1: HEADING1_SIZE, 2: HEADING2_SIZE, 3: HEADING3_SIZE,
        4: HEADING4_SIZE, 5: HEADING5_SIZE, 6: HEADING6_SIZE,
    }.items():
        _set_heading_style(doc, _lv, FONT_FAMILY_HEADING, _sz, line=240)

    # ---- 封面 ----
    for _ in range(6):
        _new_paragraph(doc, '', size=BODY_SIZE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _new_paragraph(doc, '任务说明文档', size=TITLE_SIZE, bold=True,
                   font_name=FONT_FAMILY_HEADING, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   before=200, after=100)
    _new_paragraph(doc, task_attrs.get('title', ''), size=SUBTITLE_SIZE,
                   font_name=FONT_FAMILY_HEADING, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   before=100, after=200)
    _new_paragraph(doc, '', size=BODY_SIZE)
    _new_paragraph(doc, f'项目名称：{project_name}（{project_display_id}）', size=SMALL_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _new_paragraph(doc, f'导出日期：{datetime.now().strftime("%Y年%m月%d日")}', size=SMALL_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- 正文：任务基本信息 ----
    if selected_fields:
        _add_task_info_table(doc, task_attrs, selected_fields, num_id)

    # ---- 关联需求（渲染在基本信息之后、沟通记录之前） ----
    linked_req_data = data.get('linked_requirements', [])
    req_doc_bytes_list = []  # 保存供 ZIP 打包用
    req_attachment_maps = {}  # req.id -> {(需求段, 源文件名): 导出文件名}
    if 'linked_requirements' in selected_fields and linked_req_data:
        from .routers.requirements import generate_requirement_doc_bytes, _extract_attachment_links, \
            _make_export_filename, _unique_export_name
        from .database import Requirement, RequirementCustomField, RequirementCustomValue

        all_used_names = set()  # 跨需求共享，保证整个 ZIP 内附件文件名唯一（重名加序号）
        for req_info in linked_req_data:
            # 查询完整 Requirement 对象（含 custom_values eager load）
            req = db.query(Requirement).options(
                joinedload(Requirement.custom_values).joinedload(RequirementCustomValue.field)
            ).filter(
                Requirement.id == req_info['id'],
                Requirement.project_id == proj.id
            ).first()
            if not req:
                continue

            # 构建该需求的附件导出名映射（重名序号跨需求统一计数）
            amap = {}
            for seg, fn, display in _extract_attachment_links(req.description):
                key = (seg, fn)
                if key in amap:
                    continue
                amap[key] = _unique_export_name(_make_export_filename(display, fn), all_used_names)
            req_attachment_maps[req.id] = amap

            # 调用需求页的导出方法生成 DOCX（附件链接相对 requirements/ 目录定位）
            req_bytes = generate_requirement_doc_bytes(req, proj, db, attachment_map=amap, link_prefix='')
            safe_title = _sanitize_filename(req.title)
            req_filename = f'requirements/{safe_title}_需求文档.docx'
            req_doc_bytes_list.append((safe_title, req_bytes, req_filename, req_info))

        # 在 DOCX 中渲染关联需求列表（带超链接，每行一个）
        _new_paragraph(doc, '', size=BODY_SIZE, before=80)
        _new_paragraph(doc, '关联需求：', size=BODY_SIZE, bold=True,
                       font_name=FONT_FAMILY_HEADING, before=60, after=20)
        for (safe_title, _, req_filename, req) in req_doc_bytes_list:
            display_text = f'{req["title"]}（{req.get("display_id") or ""}）'
            if display_text.endswith('（）'):
                display_text = req['title']
            p = _new_paragraph(doc, '', size=BODY_SIZE, first_line_indent=480, before=20)
            _add_hyperlink(p, display_text, req_filename, size=BODY_SIZE)

    # ---- 正文：沟通记录 ----
    h1 = _add_h1(doc, '沟通记录', num_id, 0)
    _new_paragraph(doc, f'共 {len(communications)} 条记录，按时间先后顺序排列。', size=SMALL_SIZE)

    # 内容编号实例分配器：每条记录独立实例（88, 87, 86...），避免与结构编号(99)/列表(70,200+)冲突
    _content_num_seed = [88]

    for idx, comm in enumerate(communications):
        record_num = idx + 1
        comm_time = comm['comm_at'].strftime('%Y-%m-%d %H:%M') if comm['comm_at'] else '未知时间'

        # 每条沟通记录使用独立的内容编号实例（从 1 重新开始），与结构标题的中文编号分离
        content_num_id = _setup_content_numbering(doc, _content_num_seed[0])
        _content_num_seed[0] -= 1

        # 二级标题：有主题时用"时间戳 主题"，否则保持"时间戳 记录"
        if comm['subject'].strip():
            h2 = doc.add_heading(f'{comm_time} {comm["subject"].strip()}', level=2)
        else:
            h2 = doc.add_heading(f'{comm_time} 记录', level=2)
        _apply_numbering(h2, num_id, 1)

        if not comm_minimal:
            # 元信息：时间 · 类型 · 对接人
            meta_parts = [f'沟通时间：{comm_time}']
            if comm['comm_type']:
                meta_parts.append(comm['comm_type'])
            contact_text = '、'.join(comm['contacts']) if comm['contacts'] else ''
            if contact_text:
                meta_parts.append(f'对接人：{contact_text}')
            _new_paragraph(doc, ' | '.join(meta_parts), size=SMALL_SIZE)

            # 状态变更
            if comm['old_status_id'] or comm['new_status_id']:
                if comm['old_status_name'] and comm['new_status_name']:
                    status_line = f'状态变更：{comm["old_status_name"]} → {comm["new_status_name"]}'
                elif comm['new_status_name']:
                    status_line = f'状态变更为：{comm["new_status_name"]}'
                else:
                    status_line = f'状态：{comm["old_status_name"]}（不变）'
                _new_paragraph(doc, status_line, size=SMALL_SIZE)

        # 沟通内容
        # 自动生成的状态变更文本与上方状态行重复，跳过不渲染（#8）
        if not _is_auto_status_text(comm['content'], comm['old_status_id'], comm['new_status_id']):
            if not comm_minimal:
                _new_paragraph(doc, '沟通内容：', size=BODY_SIZE, bold=True, before=80)
            if comm['content']:
                # 富文本内容（如带格式/图片/超链接）复用需求描述的 HTML→DOCX 渲染器（#4）
                if _looks_like_html(comm['content']):
                    try:
                        from .routers.requirements import _render_html_to_docx
                        # 沟通内容作为 comm_content 渲染：标题层级下移 2 级（h1->Heading3 ...），
                        # 套用本记录独立的阿拉伯编号实例（content_num_id：1 / 1.1 / (1)），
                        # 不再手填 一、/1.1/(1)
                        _render_html_to_docx(doc, comm['content'], {}, {}, img_base_dir=UPLOAD_DIR,
                                             comm_content=True, num_id=content_num_id)
                    except Exception:
                        for line in comm['content'].split('\n'):
                            if line.strip():
                                _new_paragraph(doc, line.strip(), size=BODY_SIZE, first_line_indent=480)
                else:
                    for line in comm['content'].split('\n'):
                        if line.strip():
                            _new_paragraph(doc, line.strip(), size=BODY_SIZE, first_line_indent=480)

        # ---- 附件 ----
        image_atts = [a for a in comm['attachments'] if a['is_image']]
        non_image_atts = [a for a in comm['attachments'] if not a['is_image']]

        # 非图片附件
        if non_image_atts:
            _new_paragraph(doc, '附件：', size=BODY_SIZE, bold=True, before=120)
            for att in non_image_atts:
                size_str = _format_size(att['file_size'])
                p = _new_paragraph(doc, '', size=BODY_SIZE, before=20, first_line_indent=480)
                if att['full_path']:
                    rel_path = f'attachments/{record_num}/{att["original_filename"]}'
                    _add_hyperlink(p, att['original_filename'], rel_path, size=BODY_SIZE)
                else:
                    _add_run(p, att['original_filename'], size=BODY_SIZE)
                _add_run(p, f'（{size_str}）', size=BODY_SIZE)

        # 内嵌图片
        if image_atts:
            _new_paragraph(doc, '【图片附件】', size=BODY_SIZE, bold=True, before=120)
            for att in image_atts:
                if att['full_path']:
                    try:
                        pic_p = _new_paragraph(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                               before=80, after=40)
                        pic_p.add_run().add_picture(att['full_path'], width=Inches(5.0))
                        cap_p = _new_paragraph(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER, before=20)
                        rel_path = f'attachments/{record_num}/{att["original_filename"]}'
                        _add_hyperlink(cap_p, att['original_filename'], rel_path, size=SMALL_SIZE)
                    except Exception:
                        _new_paragraph(doc, f'　　[图片加载失败：{att["original_filename"]}]',
                                       size=BODY_SIZE, color=RGBColor(180, 0, 0))
                else:
                    _new_paragraph(doc, f'　　[文件不可访问：{att["original_filename"]}]',
                                   size=BODY_SIZE, color=RGBColor(180, 0, 0))

        # 分隔线
        if idx < len(communications) - 1:
            _new_paragraph(doc, '', before=200)

    # ---- 保存 DOCX ----
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    # ======================== 生成 ZIP ========================
    zip_buffer = io.BytesIO()
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f'{task_attrs["title"]}_说明文档_{ts}.docx', docx_bytes)

        # 添加需求文档
        if req_doc_bytes_list:
            for _, req_bytes, req_filename, _ in req_doc_bytes_list:
                zf.writestr(req_filename, req_bytes)

        # 收集并添加需求附件文件
        req_file_entries = []  # (arcname, file_path)
        # 按 (需求段, 文件名) 去重，避免不同需求同名文件互相覆盖（#10）
        seen_files = set()
        for req_info in linked_req_data:
            req = db.query(Requirement).filter(
                Requirement.id == req_info['id'],
                Requirement.project_id == proj.id
            ).first()
            if not req or not req.description:
                continue
            # 同时捕获需求段与文件名，写入带段的路径，杜绝跨需求同名冲突
            file_pattern = re.compile(
                r'/uploads/' + re.escape(proj.display_id) +
                r'/requirements/([^/]+)/files/([^"\s)]+)'
            )
            for m in file_pattern.finditer(req.description):
                seg = m.group(1)
                fn = m.group(2)
                key = (seg, fn)
                if key in seen_files:
                    continue
                seen_files.add(key)
                file_path = os.path.join(
                    UPLOAD_DIR, proj.display_id, 'requirements', seg, 'files', fn
                )
                if os.path.isfile(file_path):
                    # 按导出名写入，与需求文档 DOCX 内超链接一致
                    new_name = (req_attachment_maps.get(req.id) or {}).get(key, fn)
                    arcname = f'requirements/{seg}/files/{new_name}'
                    req_file_entries.append((arcname, file_path))

        for arcname, file_path in req_file_entries:
            try:
                zf.write(file_path, arcname)
            except Exception:
                pass

        att_count = 0
        for rec_idx, comm in enumerate(communications):
            record_number = rec_idx + 1
            for att in comm['attachments']:
                if att.get('is_image'):  # 图片已内嵌到 DOCX 正文，不单独打包
                    continue
                if att['full_path']:
                    try:
                        arcname = f'attachments/{record_number}/{att["original_filename"]}'
                        zf.write(att['full_path'], arcname)
                        att_count += 1
                    except Exception:
                        pass

    return zip_buffer.getvalue(), len(communications), att_count
