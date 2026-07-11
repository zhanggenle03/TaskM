from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import Response
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func, case, select, or_
from typing import List, Optional, Dict
from collections import defaultdict
from datetime import datetime, date, timedelta
import json, os, uuid, io, urllib.parse, re, ast
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..database import (
    get_db, Project, Requirement, RequirementCustomField, RequirementCustomValue,
    RequirementStatusPool, RequirementPriorityPool,
    Task, StatusPool, touch_project, resolve_project, resolve_requirement,
    generate_requirement_display_id, UPLOAD_DIR, CONFIG_DIR,
)
from ..schemas import (
    RequirementCreate, RequirementUpdate, RequirementOut,
    RequirementCustomFieldCreate, RequirementCustomFieldUpdate, RequirementCustomFieldOut,
    RequirementCustomValueOut,
    RequirementStatusPoolCreate, RequirementStatusPoolUpdate, RequirementStatusPoolOut,
    RequirementPriorityPoolCreate, RequirementPriorityPoolUpdate, RequirementPriorityPoolOut,
    StatusDistribution, PriorityDistribution, TrendPoint,
    ProjectProgress, DashboardData, DistributionItem,
)

router = APIRouter(prefix="/projects/{project_id}/requirements", tags=["requirements"])


# 内置字段名称（不允许创建同名自定义字段，不允许删除）
BUILTIN_FIELD_NAMES = {"标题", "状态", "优先级", "创建时间", "更新时间"}
EMPTY_FILTER_VALUE = "(空)"


def _apply_empty_filter(column, values):
    """处理筛选条件中的空值标记 (空)"""
    non_empty = [v for v in values if v != EMPTY_FILTER_VALUE]
    has_empty = EMPTY_FILTER_VALUE in values
    if has_empty and non_empty:
        return or_(column.in_(non_empty), column.is_(None), column == "")
    elif has_empty:
        return or_(column.is_(None), column == "")
    else:
        return column.in_(values)


# ========== 自定义字段 CRUD ==========

@router.get("/fields", response_model=List[RequirementCustomFieldOut])
def list_custom_fields(
    project_id: str,
    show_inactive: bool = False,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    q = db.query(RequirementCustomField).filter(
        RequirementCustomField.project_id == proj.id
    )
    if not show_inactive:
        q = q.filter(RequirementCustomField.is_active == True)
    # 过滤旧版遗留的与内置字段同名的非内置字段
    q = q.filter(
        ~((RequirementCustomField.is_builtin == False) &
          RequirementCustomField.field_name.in_(list(BUILTIN_FIELD_NAMES)))
    )
    return q.order_by(RequirementCustomField.sort_order).all()


@router.post("/fields", response_model=RequirementCustomFieldOut)
def create_custom_field(
    project_id: str,
    data: RequirementCustomFieldCreate,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)

    # 不允许创建与内置字段同名的自定义字段
    if data.field_name in BUILTIN_FIELD_NAMES:
        raise HTTPException(400, f"「{data.field_name}」为内置字段，不能重复创建")

    # 自动分配 sort_order：排在最后
    max_order = db.query(func.max(RequirementCustomField.sort_order)).filter(
        RequirementCustomField.project_id == proj.id,
    ).scalar() or -1
    auto_sort = max_order + 1

    # 检查是否存在同名非活动字段，有则自动激活
    existing = db.query(RequirementCustomField).filter(
        RequirementCustomField.project_id == proj.id,
        RequirementCustomField.field_name == data.field_name,
        RequirementCustomField.is_active == False,
    ).first()
    if existing:
        existing.is_active = True
        existing.field_type = data.field_type
        existing.field_options = data.field_options
        existing.sort_order = data.sort_order if data.sort_order != 0 else auto_sort
        db.commit()
        db.refresh(existing)
        touch_project(db, proj.id)
        return existing

    field = RequirementCustomField(
        project_id=proj.id,
        field_name=data.field_name,
        field_type=data.field_type,
        field_options=data.field_options,
        sort_order=data.sort_order if data.sort_order != 0 else auto_sort,
    )
    db.add(field)
    db.commit()
    db.refresh(field)
    touch_project(db, proj.id)
    return field


@router.put("/fields/{field_id}", response_model=RequirementCustomFieldOut)
def update_custom_field(
    project_id: str,
    field_id: int,
    data: RequirementCustomFieldUpdate,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    field = db.query(RequirementCustomField).filter(
        RequirementCustomField.id == field_id,
        RequirementCustomField.project_id == proj.id
    ).first()
    if not field:
        raise HTTPException(404, "自定义字段不存在")
    # 内置字段只能修改排序和激活状态，不能改名和改类型
    if field.is_builtin:
        if data.field_name is not None and data.field_name != field.field_name:
            raise HTTPException(400, "内置字段不允许修改名称")
        if data.field_type is not None and data.field_type != field.field_type:
            raise HTTPException(400, "内置字段不允许修改类型")
        # 标题不允许停用
        if data.is_active is not None and not data.is_active and field.field_name == "标题":
            raise HTTPException(400, "标题字段不允许停用")
    if data.field_name is not None:
        field.field_name = data.field_name
    if data.field_type is not None:
        field.field_type = data.field_type
    if data.field_options is not None:
        field.field_options = data.field_options
    if data.sort_order is not None:
        field.sort_order = data.sort_order
    if data.is_active is not None:
        field.is_active = data.is_active
    db.commit()
    db.refresh(field)
    touch_project(db, proj.id)
    return field


@router.delete("/fields/{field_id}")
def delete_custom_field(
    project_id: str,
    field_id: int,
    force: bool = False,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    field = db.query(RequirementCustomField).filter(
        RequirementCustomField.id == field_id,
        RequirementCustomField.project_id == proj.id
    ).first()
    if not field:
        raise HTTPException(404, "自定义字段不存在")
    if field.is_builtin:
        raise HTTPException(400, "内置字段不允许删除，只能停用")
    if force:
        # 批量更新受影响需求的 updated_at（自定义值将被级联删除）
        req_ids = db.query(RequirementCustomValue.requirement_id).filter(
            RequirementCustomValue.field_id == field_id
        ).distinct().all()
        if req_ids:
            affected = [r[0] for r in req_ids]
            db.query(Requirement).filter(Requirement.id.in_(affected)).update(
                {"updated_at": datetime.now()}, synchronize_session=False
            )
        db.delete(field)
    else:
        field.is_active = False
    db.commit()
    touch_project(db, proj.id)
    return {"message": "ok"}


@router.get("/fields/{field_id}/existing-values")
def get_custom_field_existing_values(
    project_id: str,
    field_id: int,
    db: Session = Depends(get_db),
):
    """获取自定义字段在需求列表中已存在的值列表（去重），供自动填充选项使用。"""
    proj = resolve_project(db, project_id)
    field = db.query(RequirementCustomField).filter(
        RequirementCustomField.id == field_id,
        RequirementCustomField.project_id == proj.id
    ).first()
    if not field:
        raise HTTPException(404, "自定义字段不存在")

    # 查询该字段所有非空的值
    rows = db.query(RequirementCustomValue.value).filter(
        RequirementCustomValue.field_id == field_id,
        RequirementCustomValue.value.isnot(None),
        RequirementCustomValue.value != "",
    ).distinct().all()

    values_set: set[str] = set()
    for (val,) in rows:
        if not val:
            continue
        if field.field_type == "multi_dropdown":
            # multi_dropdown 值存为 Python 列表字符串，如 "['标签A', '标签B']"
            raw = val.strip()
            parsed_items = []
            if raw.startswith("[") and raw.endswith("]"):
                try:
                    parsed = ast.literal_eval(raw)
                    if isinstance(parsed, (list, tuple)):
                        parsed_items = [str(item).strip() for item in parsed if str(item).strip()]
                except (ValueError, SyntaxError):
                    pass
            if not parsed_items:
                parsed_items = [s.strip() for s in raw.split(",") if s.strip()]
            for item in parsed_items:
                if item:
                    values_set.add(item)
        else:
            values_set.add(val.strip())

    # 返回排序后的列表
    result = sorted(values_set)
    return result


# ========== 筛选面板统计 ==========

@router.get("/filter-stats")
def filter_stats(
    project_id: str,
    column_filters: Optional[str] = None,  # JSON: {prop: [values]} 跨列联动
    fuzzy_filters: Optional[str] = None,   # JSON: {prop: {text, mode: 'include'|'exclude'}} 模糊筛选
    db: Session = Depends(get_db),
):
    """返回所有可筛选列的独立值及其出现次数（全量数据，不分页）
    若传入 column_filters，则按其他列筛选后再统计（用于跨列联动）。
    若传入 fuzzy_filters，则按模糊筛选后再统计（用于模糊搜索联动）。
    """
    proj = resolve_project(db, project_id)
    req_q = db.query(Requirement).filter(
        Requirement.project_id == proj.id
    )

    # 应用跨列筛选（排除当前列的话由前端控制）
    if column_filters:
        try:
            cf_json = json.loads(column_filters)
        except (json.JSONDecodeError, TypeError):
            cf_json = {}
        for prop, values in cf_json.items():
            if not values:
                continue
            if prop == "status":
                req_q = req_q.filter(_apply_empty_filter(Requirement.status, values))
            elif prop == "priority":
                req_q = req_q.filter(_apply_empty_filter(Requirement.priority, values))
            elif prop == "display_id":
                req_q = req_q.filter(_apply_empty_filter(Requirement.display_id, values))
            elif prop == "title":
                req_q = req_q.filter(_apply_empty_filter(Requirement.title, values))
            elif prop == "created_at":
                req_q = req_q.filter(_apply_empty_filter(Requirement.created_at, values))
            elif prop == "updated_at":
                req_q = req_q.filter(_apply_empty_filter(Requirement.updated_at, values))
            elif prop.startswith("cf_"):
                try:
                    fid = int(prop[3:])
                except ValueError:
                    continue
                cf = db.query(RequirementCustomField).filter(
                    RequirementCustomField.id == fid,
                    RequirementCustomField.project_id == proj.id,
                ).first()
                if not cf:
                    continue
                non_empty = [v for v in values if v != EMPTY_FILTER_VALUE]
                has_empty = EMPTY_FILTER_VALUE in values
                if cf.field_type in ("date", "datetime"):
                    subq = db.query(RequirementCustomValue.requirement_id).filter(
                        RequirementCustomValue.field_id == fid
                    )
                    like_conds = [RequirementCustomValue.value.like(f"{v}%") for v in non_empty if v]

                    if has_empty and like_conds:
                        req_q = req_q.filter(or_(
                            Requirement.id.in_(subq.filter(or_(*like_conds))),
                            ~Requirement.id.in_(subq),
                        ))
                    elif has_empty:
                        req_q = req_q.filter(~Requirement.id.in_(subq))
                    elif like_conds:
                        req_q = req_q.filter(Requirement.id.in_(subq.filter(or_(*like_conds))))
                else:
                    subq = db.query(RequirementCustomValue.requirement_id).filter(
                        RequirementCustomValue.field_id == fid,
                    )
                    if has_empty and non_empty:
                        req_q = req_q.filter(or_(
                            Requirement.id.in_(subq.filter(RequirementCustomValue.value.in_(non_empty))),
                            ~Requirement.id.in_(subq),
                        ))
                    elif has_empty:
                        req_q = req_q.filter(~Requirement.id.in_(subq))
                    else:
                        req_q = req_q.filter(Requirement.id.in_(subq.filter(RequirementCustomValue.value.in_(values))))

    # 模糊筛选（fuzzy_filters）
    if fuzzy_filters:
        try:
            fuzzy_json = json.loads(fuzzy_filters)
        except (json.JSONDecodeError, TypeError):
            fuzzy_json = {}
        for prop, f in fuzzy_json.items():
            text = (f.get("text") or "").strip()
            if not text:
                continue
            mode = f.get("mode", "include")
            keywords = [k for k in text.split() if k]
            if not keywords:
                continue
            if prop == "status":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    req_q = req_q.filter(Requirement.status.like(like_val) if mode == 'include' else ~Requirement.status.like(like_val))
            elif prop == "priority":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    req_q = req_q.filter(Requirement.priority.like(like_val) if mode == 'include' else ~Requirement.priority.like(like_val))
            elif prop == "display_id":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    req_q = req_q.filter(Requirement.display_id.like(like_val) if mode == 'include' else ~Requirement.display_id.like(like_val))
            elif prop == "title":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    req_q = req_q.filter(Requirement.title.like(like_val) if mode == 'include' else ~Requirement.title.like(like_val))
            elif prop == "created_at":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    req_q = req_q.filter(Requirement.created_at.like(like_val) if mode == 'include' else ~Requirement.created_at.like(like_val))
            elif prop == "updated_at":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    req_q = req_q.filter(Requirement.updated_at.like(like_val) if mode == 'include' else ~Requirement.updated_at.like(like_val))
            elif prop.startswith("cf_"):
                try:
                    fid = int(prop[3:])
                except ValueError:
                    continue
                for kw in keywords:
                    like_val = f"%{kw}%"
                    subq = db.query(RequirementCustomValue.requirement_id).filter(
                        RequirementCustomValue.field_id == fid,
                        RequirementCustomValue.value.like(like_val),
                    )
                    if mode == 'include':
                        req_q = req_q.filter(Requirement.id.in_(subq))
                    else:
                        req_q = req_q.filter(~Requirement.id.in_(subq))

    reqs = req_q.all()

    # 内置列（转换为前端中文标签）
    STATUS_LABELS = {'todo': '待处理', 'in_progress': '进行中', 'done': '已完成', 'cancelled': '已取消'}
    PRIORITY_LABELS = {'low': '低', 'normal': '普通', 'high': '高', 'urgent': '紧急'}
    status_counter = {}
    priority_counter = {}
    display_id_counter = {}
    title_counter = {}
    created_counter = {}
    updated_counter = {}
    for r in reqs:
        status_counter[STATUS_LABELS.get(r.status, r.status) if r.status else EMPTY_FILTER_VALUE] = status_counter.get(STATUS_LABELS.get(r.status, r.status) if r.status else EMPTY_FILTER_VALUE, 0) + 1
        priority_counter[PRIORITY_LABELS.get(r.priority, r.priority) if r.priority else EMPTY_FILTER_VALUE] = priority_counter.get(PRIORITY_LABELS.get(r.priority, r.priority) if r.priority else EMPTY_FILTER_VALUE, 0) + 1
        display_id_counter[r.display_id or EMPTY_FILTER_VALUE] = display_id_counter.get(r.display_id or EMPTY_FILTER_VALUE, 0) + 1
        title_counter[r.title or EMPTY_FILTER_VALUE] = title_counter.get(r.title or EMPTY_FILTER_VALUE, 0) + 1
        if r.created_at:
            dt = r.created_at[:19] if isinstance(r.created_at, str) else r.created_at.strftime("%Y-%m-%d %H:%M:%S")
            created_counter[dt] = created_counter.get(dt, 0) + 1
        else:
            created_counter[EMPTY_FILTER_VALUE] = created_counter.get(EMPTY_FILTER_VALUE, 0) + 1
        if r.updated_at:
            dt = r.updated_at[:19] if isinstance(r.updated_at, str) else r.updated_at.strftime("%Y-%m-%d %H:%M:%S")
            updated_counter[dt] = updated_counter.get(dt, 0) + 1
        else:
            updated_counter[EMPTY_FILTER_VALUE] = updated_counter.get(EMPTY_FILTER_VALUE, 0) + 1

    def fmt(counter):
        return [{"value": k, "count": v} for k, v in sorted(counter.items(), key=lambda x: -x[1])]

    result = {
        "status": fmt(status_counter),
        "priority": fmt(priority_counter),
        "display_id": fmt(display_id_counter),
        "title": fmt(title_counter),
        "created_at": fmt(created_counter),
        "updated_at": fmt(updated_counter),
    }

    # 自定义字段
    custom_fields = db.query(RequirementCustomField).filter(
        RequirementCustomField.project_id == proj.id,
        RequirementCustomField.is_active == True,
    ).all()

    for cf in custom_fields:
        cv_counter = {}
        for r in reqs:
            cv = db.query(RequirementCustomValue).filter(
                RequirementCustomValue.requirement_id == r.id,
                RequirementCustomValue.field_id == cf.id,
            ).first()
            val = (cv.value if cv and cv.value else None) or EMPTY_FILTER_VALUE
            cv_counter[val] = cv_counter.get(val, 0) + 1
        result[f"cf_{cf.id}"] = fmt(cv_counter)

    return result


# ========== 大屏数据统计 API ==========

@router.get("/stats/dashboard", response_model=DashboardData)
def dashboard_stats(
    project_id: str,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)

    # 1. 需求状态分布
    status_counts = db.query(
        Requirement.status, func.count(Requirement.id)
    ).filter(
        Requirement.project_id == proj.id
    ).group_by(Requirement.status).all()

    status_label_map = {
        "todo": "待处理",
        "in_progress": "进行中",
        "done": "已完成",
        "cancelled": "已取消",
    }
    status_distribution = [
        StatusDistribution(name=status_label_map.get(s, s), value=c)
        for s, c in status_counts
    ]

    # 2. 需求优先级分布
    priority_counts = db.query(
        Requirement.priority, func.count(Requirement.id)
    ).filter(
        Requirement.project_id == proj.id
    ).group_by(Requirement.priority).all()

    priority_label_map = {
        "low": "低",
        "normal": "普通",
        "high": "高",
        "urgent": "紧急",
    }
    priority_distribution = [
        PriorityDistribution(name=priority_label_map.get(p, p), value=c)
        for p, c in priority_counts
    ]

    # 3. 项目进度概览（任务状态分布）
    status_pools = db.query(StatusPool).filter(
        StatusPool.project_id == proj.id,
        StatusPool.is_active == True,
    ).all()
    project_progress = []
    for sp in status_pools:
        count = db.query(Task).filter(
            Task.project_id == proj.id,
            Task.status_id == sp.id,
        ).count()
        if count > 0:
            project_progress.append(ProjectProgress(
                name=sp.name, value=count, color=sp.color
            ))

    # 4. 需求趋势（按周统计创建量，最近10周）
    today = date.today()
    trend_data = []
    for i in range(9, -1, -1):
        week_start = today - timedelta(days=today.weekday()) - timedelta(weeks=i)
        week_end = week_start + timedelta(days=6)
        count = db.query(func.count(Requirement.id)).filter(
            Requirement.project_id == proj.id,
            Requirement.created_at >= datetime.combine(week_start, datetime.min.time()),
            Requirement.created_at <= datetime.combine(week_end, datetime.max.time()),
        ).scalar() or 0
        trend_data.append(TrendPoint(
            date=week_start.strftime("%m/%d"),
            count=count,
        ))

    # 5. 自定义字段分布（dropdown/multi_dropdown 类型）
    custom_fields = db.query(RequirementCustomField).filter(
        RequirementCustomField.project_id == proj.id,
        RequirementCustomField.is_active == True,
        RequirementCustomField.field_type.in_(["dropdown", "multi_dropdown"]),
        # 排除与内置字段同名的自定义字段，避免下拉框重复
        ~RequirementCustomField.field_name.in_(list(BUILTIN_FIELD_NAMES)),
    ).all()

    extra_distributions: dict[str, list[DistributionItem]] = {}
    available_chart_fields: list[dict] = [
        {"key": "priority", "label": "优先级"},
        {"key": "status", "label": "状态"},
    ]

    # 获取所有需求的 custom_values，按 field_id 分组
    all_cvs = db.query(RequirementCustomValue).options(
        joinedload(RequirementCustomValue.field)
    ).filter(
        RequirementCustomValue.field_id.in_([f.id for f in custom_fields]),
    ).all()
    cv_by_field: dict[int, list[str]] = defaultdict(list)
    for cv in all_cvs:
        values = cv.value or ""
        if cv.field.field_type == "multi_dropdown":
            # multi_dropdown 值存为 Python 列表字符串，如 "['标签A', '标签B']"
            raw = values.strip()
            parsed_items = []
            # 优先用 ast.literal_eval 解析（最可靠）
            if raw.startswith("[") and raw.endswith("]"):
                try:
                    parsed = ast.literal_eval(raw)
                    if isinstance(parsed, (list, tuple)):
                        parsed_items = [str(item).strip() for item in parsed if str(item).strip()]
                except (ValueError, SyntaxError):
                    pass
            # 降级：逗号拆分（兼容其他存储格式）
            if not parsed_items:
                parsed_items = [s.strip() for s in raw.split(",") if s.strip()]
            for item in parsed_items:
                cv_by_field[cv.field_id].append(item)
        else:
            cv_by_field[cv.field_id].append(values)

    for f in custom_fields:
        available_chart_fields.append({"key": f"cf_{f.id}", "label": f.field_name})
        counter: dict[str, int] = defaultdict(int)
        for val in cv_by_field.get(f.id, []):
            if val.strip():
                counter[val.strip()] += 1
        extra_distributions[f"cf_{f.id}"] = [
            DistributionItem(name=k, value=v) for k, v in counter.items()
        ]

    return DashboardData(
        status_distribution=status_distribution,
        priority_distribution=priority_distribution,
        project_progress=project_progress,
        trend=trend_data,
        extra_distributions=extra_distributions,
        available_chart_fields=available_chart_fields,
    )


# ----  Dashboard 看板配置（kanban.json） ----

def _dashboard_kanban_config_path(proj):
    d = os.path.join(CONFIG_DIR, proj.display_id)
    return os.path.join(d, "kanban.json")


@router.get("/kanban-config")
def get_req_kanban_config(project_id: str, db: Session = Depends(get_db)):
    proj = resolve_project(db, project_id)
    p = _dashboard_kanban_config_path(proj)
    if not os.path.isfile(p):
        return {}
    try:
        with open(p, encoding="utf-8") as f:
            return json.load(f)
    except:
        return {}


@router.put("/kanban-config")
def put_req_kanban_config(project_id: str, body: dict, db: Session = Depends(get_db)):
    proj = resolve_project(db, project_id)
    p = _dashboard_kanban_config_path(proj)
    os.makedirs(os.path.dirname(p), exist_ok=True)
    # 读取现有配置，合并新数据，保留已有 key（如 task kanban 的列配置）
    existing = {}
    if os.path.isfile(p):
        try:
            with open(p, encoding="utf-8") as f:
                existing = json.load(f)
        except:
            pass
    existing.update(body)
    with open(p, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)
    return {"ok": True}


# ========== 辅助函数 ==========

def _get_requirement_with_values(db: Session, req_id: int) -> RequirementOut:
    req = db.query(Requirement).options(
        joinedload(Requirement.custom_values).joinedload(RequirementCustomValue.field)
    ).filter(Requirement.id == req_id).first()
    return _format_requirement(req) if req else None


def _format_requirement(req: Requirement) -> RequirementOut:
    vals = []
    for cv in req.custom_values:
        if cv.field and cv.field.is_builtin:
            continue
        vals.append(RequirementCustomValueOut(
            field_id=cv.field_id,
            field_name=cv.field.field_name if cv.field else "",
            field_type=cv.field.field_type if cv.field else "",
            value=cv.value,
        ))
    return RequirementOut(
        id=req.id,
        project_id=req.project_id,
        display_id=req.display_id,
        title=req.title,
        description=req.description or "",
        priority=req.priority,
        status=req.status,
        created_at=req.created_at,
        updated_at=req.updated_at,
        custom_values=vals,
    )


# ========== 需求状态池 ==========

def _sync_builtin_field_options(proj, db):
    """同步内置字段的 field_options 与对应池的值"""
    # 状态
    status_names = [s.name for s in db.query(RequirementStatusPool).filter(
        RequirementStatusPool.project_id == proj.id,
        RequirementStatusPool.is_active == True,
    ).order_by(RequirementStatusPool.sort_order).all()]
    status_field = db.query(RequirementCustomField).filter(
        RequirementCustomField.project_id == proj.id,
        RequirementCustomField.field_name == "状态",
        RequirementCustomField.is_builtin == True,
    ).first()
    if status_field:
        status_field.field_options = "\n".join(status_names) if status_names else ""

    # 优先级
    priority_names = [p.name for p in db.query(RequirementPriorityPool).filter(
        RequirementPriorityPool.project_id == proj.id,
        RequirementPriorityPool.is_active == True,
    ).order_by(RequirementPriorityPool.sort_order).all()]
    priority_field = db.query(RequirementCustomField).filter(
        RequirementCustomField.project_id == proj.id,
        RequirementCustomField.field_name == "优先级",
        RequirementCustomField.is_builtin == True,
    ).first()
    if priority_field:
        priority_field.field_options = "\n".join(priority_names) if priority_names else ""


@router.get("/status-pools", response_model=List[RequirementStatusPoolOut])
def list_status_pools(
    project_id: str,
    show_inactive: bool = False,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    q = db.query(RequirementStatusPool).filter(RequirementStatusPool.project_id == proj.id)
    if not show_inactive:
        q = q.filter(RequirementStatusPool.is_active == True)
    return q.order_by(RequirementStatusPool.sort_order).all()


@router.post("/status-pools", response_model=RequirementStatusPoolOut)
def create_status_pool(
    project_id: str,
    data: RequirementStatusPoolCreate,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    existing = db.query(RequirementStatusPool).filter(
        RequirementStatusPool.project_id == proj.id,
        RequirementStatusPool.name == data.name,
        RequirementStatusPool.is_active == False,
    ).first()
    if existing:
        existing.is_active = True
        existing.color = data.color
        existing.sort_order = data.sort_order
        existing.is_default = data.is_default
        _sync_builtin_field_options(proj, db)
        db.commit()
        db.refresh(existing)
        return existing
    if data.is_default:
        db.query(RequirementStatusPool).filter(
            RequirementStatusPool.project_id == proj.id,
            RequirementStatusPool.is_default == True,
        ).update({"is_default": False})
    pool = RequirementStatusPool(project_id=proj.id, **data.model_dump())
    db.add(pool)
    _sync_builtin_field_options(proj, db)
    db.commit()
    db.refresh(pool)
    touch_project(db, proj.id)
    return pool


@router.put("/status-pools/{pool_id}", response_model=RequirementStatusPoolOut)
def update_status_pool(project_id: str, pool_id: int, data: RequirementStatusPoolUpdate, db: Session = Depends(get_db)):
    proj = resolve_project(db, project_id)
    pool = db.query(RequirementStatusPool).filter(
        RequirementStatusPool.id == pool_id, RequirementStatusPool.project_id == proj.id
    ).first()
    if not pool:
        raise HTTPException(404, "状态不存在")
    if data.is_default and not pool.is_default:
        db.query(RequirementStatusPool).filter(
            RequirementStatusPool.project_id == proj.id, RequirementStatusPool.is_default == True
        ).update({"is_default": False})
    update_data = data.model_dump(exclude_unset=True)
    for k, v in update_data.items():
        setattr(pool, k, v)
    _sync_builtin_field_options(proj, db)
    db.commit()
    db.refresh(pool)
    return pool


@router.delete("/status-pools/{pool_id}")
def delete_status_pool(
    project_id: str, pool_id: int,
    force: bool = False,
    confirmed: bool = False,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    pool = db.query(RequirementStatusPool).filter(
        RequirementStatusPool.id == pool_id, RequirementStatusPool.project_id == proj.id
    ).first()
    if not pool:
        raise HTTPException(404, "状态不存在")

    # 默认值不允许停用或删除
    if pool.is_default:
        raise HTTPException(400, "默认值不允许停用或删除，请先设置其他状态为默认值")

    # 硬删除：检查引用数量
    if force:
        # 映射 name → requirement status value
        name_to_value = {"待处理": "todo", "进行中": "in_progress", "已完成": "done", "已取消": "cancelled"}
        status_value = name_to_value.get(pool.name, pool.name)
        ref_count = db.query(func.count(Requirement.id)).filter(
            Requirement.project_id == proj.id,
            Requirement.status == status_value,
        ).scalar() or 0

        if ref_count > 0 and not confirmed:
            raise HTTPException(409, detail={
                "message": "该状态被引用",
                "refs": {"需求": ref_count},
                "refs_cleaned": None,
            })

        if confirmed:
            # 查找默认状态
            default_status = db.query(RequirementStatusPool).filter(
                RequirementStatusPool.project_id == proj.id,
                RequirementStatusPool.is_default == True,
                RequirementStatusPool.id != pool_id,
            ).first()
            default_value = None
            if default_status:
                default_value = name_to_value.get(default_status.name, default_status.name)
            # 重置引用
            db.query(Requirement).filter(
                Requirement.project_id == proj.id,
                Requirement.status == status_value,
            ).update({"status": default_value or "todo"}, synchronize_session=False)

        db.delete(pool)
        _sync_builtin_field_options(proj, db)
        db.commit()
        return {"message": "ok", "refs_cleaned": {"需求": ref_count} if ref_count > 0 else {}}

    pool.is_active = False
    _sync_builtin_field_options(proj, db)
    db.commit()
    return {"message": "ok"}


# ========== 需求优先级池 ==========

@router.get("/priority-pools", response_model=List[RequirementPriorityPoolOut])
def list_priority_pools(
    project_id: str,
    show_inactive: bool = False,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    q = db.query(RequirementPriorityPool).filter(RequirementPriorityPool.project_id == proj.id)
    if not show_inactive:
        q = q.filter(RequirementPriorityPool.is_active == True)
    return q.order_by(RequirementPriorityPool.sort_order).all()


@router.post("/priority-pools", response_model=RequirementPriorityPoolOut)
def create_priority_pool(
    project_id: str,
    data: RequirementPriorityPoolCreate,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    existing = db.query(RequirementPriorityPool).filter(
        RequirementPriorityPool.project_id == proj.id,
        RequirementPriorityPool.name == data.name,
        RequirementPriorityPool.is_active == False,
    ).first()
    if existing:
        existing.is_active = True
        existing.color = data.color
        existing.sort_order = data.sort_order
        existing.is_default = data.is_default
        db.commit()
        db.refresh(existing)
        return existing
    if data.is_default:
        db.query(RequirementPriorityPool).filter(
            RequirementPriorityPool.project_id == proj.id,
            RequirementPriorityPool.is_default == True,
        ).update({"is_default": False})
    pool = RequirementPriorityPool(project_id=proj.id, **data.model_dump())
    db.add(pool)
    db.commit()
    db.refresh(pool)
    touch_project(db, proj.id)
    return pool


@router.put("/priority-pools/{pool_id}", response_model=RequirementPriorityPoolOut)
def update_priority_pool(project_id: str, pool_id: int, data: RequirementPriorityPoolUpdate, db: Session = Depends(get_db)):
    proj = resolve_project(db, project_id)
    pool = db.query(RequirementPriorityPool).filter(
        RequirementPriorityPool.id == pool_id, RequirementPriorityPool.project_id == proj.id
    ).first()
    if not pool:
        raise HTTPException(404, "优先级不存在")
    if data.is_default and not pool.is_default:
        db.query(RequirementPriorityPool).filter(
            RequirementPriorityPool.project_id == proj.id, RequirementPriorityPool.is_default == True
        ).update({"is_default": False})
    update_data = data.model_dump(exclude_unset=True)
    for k, v in update_data.items():
        setattr(pool, k, v)
    _sync_builtin_field_options(proj, db)
    db.commit()
    db.refresh(pool)
    return pool


@router.delete("/priority-pools/{pool_id}")
def delete_priority_pool(
    project_id: str, pool_id: int,
    force: bool = False,
    confirmed: bool = False,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    pool = db.query(RequirementPriorityPool).filter(
        RequirementPriorityPool.id == pool_id, RequirementPriorityPool.project_id == proj.id
    ).first()
    if not pool:
        raise HTTPException(404, "优先级不存在")

    # 默认值不允许停用或删除
    if pool.is_default:
        raise HTTPException(400, "默认值不允许停用或删除，请先设置其他优先级为默认值")

    # 硬删除：检查引用数量
    if force:
        name_to_value = {"低": "low", "普通": "normal", "高": "high", "紧急": "urgent"}
        priority_value = name_to_value.get(pool.name, pool.name)
        ref_count = db.query(func.count(Requirement.id)).filter(
            Requirement.project_id == proj.id,
            Requirement.priority == priority_value,
        ).scalar() or 0

        if ref_count > 0 and not confirmed:
            raise HTTPException(409, detail={
                "message": "该优先级被引用",
                "refs": {"需求": ref_count},
                "refs_cleaned": None,
            })

        if confirmed:
            default_priority = db.query(RequirementPriorityPool).filter(
                RequirementPriorityPool.project_id == proj.id,
                RequirementPriorityPool.is_default == True,
                RequirementPriorityPool.id != pool_id,
            ).first()
            default_value = None
            if default_priority:
                default_value = name_to_value.get(default_priority.name, default_priority.name)
            db.query(Requirement).filter(
                Requirement.project_id == proj.id,
                Requirement.priority == priority_value,
            ).update({"priority": default_value or "normal"}, synchronize_session=False)

        db.delete(pool)
        _sync_builtin_field_options(proj, db)
        db.commit()
        return {"message": "ok", "refs_cleaned": {"需求": ref_count} if ref_count > 0 else {}}

    pool.is_active = False
    _sync_builtin_field_options(proj, db)
    db.commit()
    return {"message": "ok"}


# ========== 需求 CRUD ==========

@router.get("")
def list_requirements(
    project_id: str,
    status: Optional[str] = None,
    priority: Optional[str] = None,
    search: Optional[str] = None,
    sort_by: Optional[str] = None,
    sort_order: Optional[str] = None,
    status_order: Optional[str] = None,
    priority_order: Optional[str] = None,
    column_filters: Optional[str] = None,  # JSON: {prop: [values]}
    fuzzy_filters: Optional[str] = None,   # JSON: {prop: {text, mode: 'include'|'exclude'}}
    page: int = 1,
    page_size: int = 50,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    q = db.query(Requirement).options(
        joinedload(Requirement.custom_values).joinedload(RequirementCustomValue.field)
    ).filter(Requirement.project_id == proj.id)

    if status:
        statuses = [s.strip() for s in status.split(",") if s.strip()]
        if statuses:
            q = q.filter(Requirement.status.in_(statuses))
    if priority:
        priorities = [p.strip() for p in priority.split(",") if p.strip()]
        if priorities:
            q = q.filter(Requirement.priority.in_(priorities))
    if search:
        like = f"%{search}%"
        q = q.filter(
            Requirement.title.like(like)
        )

    # 列筛选（column_filters JSON）
    if column_filters:
        try:
            cf_json = json.loads(column_filters)
        except (json.JSONDecodeError, TypeError):
            cf_json = {}
        for prop, values in cf_json.items():
            if not values:
                continue
            if prop == "status":
                q = q.filter(_apply_empty_filter(Requirement.status, values))
            elif prop == "priority":
                q = q.filter(_apply_empty_filter(Requirement.priority, values))
            elif prop == "display_id":
                q = q.filter(_apply_empty_filter(Requirement.display_id, values))
            elif prop == "title":
                q = q.filter(_apply_empty_filter(Requirement.title, values))
            elif prop == "created_at":
                q = q.filter(_apply_empty_filter(Requirement.created_at, values))
            elif prop == "updated_at":
                q = q.filter(_apply_empty_filter(Requirement.updated_at, values))
            elif prop.startswith("cf_"):
                try:
                    fid = int(prop[3:])
                except ValueError:
                    continue
                cf = db.query(RequirementCustomField).filter(
                    RequirementCustomField.id == fid,
                    RequirementCustomField.project_id == proj.id,
                ).first()
                if not cf:
                    continue
                non_empty = [v for v in values if v != EMPTY_FILTER_VALUE]
                has_empty = EMPTY_FILTER_VALUE in values
                if cf.field_type in ("date", "datetime"):
                    subq = db.query(RequirementCustomValue.requirement_id).filter(
                        RequirementCustomValue.field_id == fid
                    )
                    like_conds = [RequirementCustomValue.value.like(f"{v}%") for v in non_empty if v]

                    if has_empty and like_conds:
                        q = q.filter(or_(
                            Requirement.id.in_(subq.filter(or_(*like_conds))),
                            ~Requirement.id.in_(subq),
                        ))
                    elif has_empty:
                        q = q.filter(~Requirement.id.in_(subq))
                    elif like_conds:
                        q = q.filter(Requirement.id.in_(subq.filter(or_(*like_conds))))
                else:
                    subq = db.query(RequirementCustomValue.requirement_id).filter(
                        RequirementCustomValue.field_id == fid,
                    )
                    non_empty_vals = [v for v in values if v != EMPTY_FILTER_VALUE]
                    if has_empty and non_empty_vals:
                        q = q.filter(or_(
                            Requirement.id.in_(subq.filter(RequirementCustomValue.value.in_(non_empty_vals))),
                            ~Requirement.id.in_(subq),
                        ))
                    elif has_empty:
                        q = q.filter(~Requirement.id.in_(subq))
                    else:
                        q = q.filter(Requirement.id.in_(subq.filter(RequirementCustomValue.value.in_(values))))

    # 模糊筛选（fuzzy_filters）
    if fuzzy_filters:
        try:
            fuzzy_json = json.loads(fuzzy_filters)
        except (json.JSONDecodeError, TypeError):
            fuzzy_json = {}
        for prop, f in fuzzy_json.items():
            text = (f.get("text") or "").strip()
            if not text:
                continue
            mode = f.get("mode", "include")
            # 空格分隔多关键词，AND 逻辑
            keywords = [k for k in text.split() if k]
            if not keywords:
                continue
            if prop == "status":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    q = q.filter(Requirement.status.like(like_val) if mode == 'include' else ~Requirement.status.like(like_val))
            elif prop == "priority":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    q = q.filter(Requirement.priority.like(like_val) if mode == 'include' else ~Requirement.priority.like(like_val))
            elif prop == "display_id":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    q = q.filter(Requirement.display_id.like(like_val) if mode == 'include' else ~Requirement.display_id.like(like_val))
            elif prop == "title":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    q = q.filter(Requirement.title.like(like_val) if mode == 'include' else ~Requirement.title.like(like_val))
            elif prop == "created_at":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    q = q.filter(Requirement.created_at.like(like_val) if mode == 'include' else ~Requirement.created_at.like(like_val))
            elif prop == "updated_at":
                for kw in keywords:
                    like_val = f"%{kw}%"
                    q = q.filter(Requirement.updated_at.like(like_val) if mode == 'include' else ~Requirement.updated_at.like(like_val))
            elif prop.startswith("cf_"):
                try:
                    fid = int(prop[3:])
                except ValueError:
                    continue
                for kw in keywords:
                    like_val = f"%{kw}%"
                    subq = db.query(RequirementCustomValue.requirement_id).filter(
                        RequirementCustomValue.field_id == fid,
                        RequirementCustomValue.value.like(like_val),
                    )
                    if mode == 'include':
                        q = q.filter(Requirement.id.in_(subq))
                    else:
                        q = q.filter(~Requirement.id.in_(subq))

    # 多列排序：接收逗号分隔的 sort_by 和 sort_order
    # 如 sort_by=priority,status&sort_order=asc,desc → 先按优先级升序，再按状态降序
    sort_map = {
        "display_id": Requirement.display_id,
        "updated_at": Requirement.updated_at,
        "created_at": Requirement.created_at,
        "title": Requirement.title,
        "priority": Requirement.priority,
        "status": Requirement.status,
    }

    # 解析 status_order / priority_order 为值→索引映射（用于 CASE 表达式）
    status_idx = {}
    if status_order:
        for i, v in enumerate(status_order.split(",")):
            status_idx[v.strip()] = i
    priority_idx = {}
    if priority_order:
        for i, v in enumerate(priority_order.split(",")):
            priority_idx[v.strip()] = i

    order_clauses = []
    if sort_by and sort_order:
        sort_by_list = [s.strip() for s in sort_by.split(",") if s.strip()]
        sort_order_list = [s.strip() for s in sort_order.split(",") if s.strip()]

        for i, col_name in enumerate(sort_by_list):
            order_dir = sort_order_list[i] if i < len(sort_order_list) else "asc"

            # 自定义字段排序：使用关联子查询取字段值
            if col_name.startswith("cf_"):
                try:
                    field_id = int(col_name.replace("cf_", ""))
                except ValueError:
                    continue
                cv_subq = (
                    select(RequirementCustomValue.value)
                    .where(
                        RequirementCustomValue.requirement_id == Requirement.id,
                        RequirementCustomValue.field_id == field_id,
                    )
                    .correlate(Requirement)
                    .scalar_subquery()
                )
                order_clauses.append(cv_subq.desc() if order_dir == "desc" else cv_subq.asc())
                continue

            col = sort_map.get(col_name)
            if col is None:
                continue

            # 对 status 和 priority 使用 CASE 表达式实现池顺序排序
            if col_name == "status" and status_idx:
                col = case(status_idx, value=Requirement.status)
            elif col_name == "priority" and priority_idx:
                col = case(priority_idx, value=Requirement.priority)
            order_clauses.append(col.desc() if order_dir == "desc" else col.asc())

    # 无活跃排序时默认按更新时间降序
    if not order_clauses:
        order_clauses.append(Requirement.updated_at.desc())

    # 总记录数（分页用）
    total = q.count()

    # 分页
    page = max(1, page)
    page_size = min(max(1, page_size), 9999)
    results = q.order_by(*order_clauses).offset((page - 1) * page_size).limit(page_size).all()

    # 组装 custom_values 为嵌套结构
    out = []
    for r in results:
        vals = []
        for cv in r.custom_values:
            if cv.field and cv.field.is_builtin:
                continue
            vals.append(RequirementCustomValueOut(
                field_id=cv.field_id,
                field_name=cv.field.field_name if cv.field else "",
                field_type=cv.field.field_type if cv.field else "",
                value=cv.value,
            ))
        out.append(RequirementOut(
            id=r.id,
            project_id=r.project_id,
            display_id=r.display_id,
            title=r.title,
            priority=r.priority,
            status=r.status,
            created_at=r.created_at,
            updated_at=r.updated_at,
            custom_values=vals,
        ))
    return {"items": out, "total": total}



@router.post("", response_model=RequirementOut)
def create_requirement(
    project_id: str,
    data: RequirementCreate,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    req = Requirement(
        project_id=proj.id,
        display_id=generate_requirement_display_id(db, proj),
        title=data.title,
        priority=data.priority,
        status=data.status,
    )
    db.add(req)
    db.flush()  # 获取 req.id

    # 保存自定义字段值
    if data.custom_values:
        for field_id_str, value in data.custom_values.items():
            try:
                fid = int(field_id_str)
            except ValueError:
                continue
            if value is not None and value != "":
                cv = RequirementCustomValue(
                    requirement_id=req.id,
                    field_id=fid,
                    value=str(value),
                )
                db.add(cv)

    db.commit()
    db.refresh(req)
    touch_project(db, proj.id)

    # 重新加载含 custom_values 的完整对象
    return _get_requirement_with_values(db, req.id)


# ========== 需求列宽持久化 ==========

COL_WIDTHS_FILE = "column_widths.json"


def _col_widths_path(proj):
    """返回列宽 JSON 文件的完整路径"""
    d = os.path.join(CONFIG_DIR, proj.display_id)
    return os.path.join(d, COL_WIDTHS_FILE)


@router.get("/column-widths")
def get_column_widths(project_id: str, db: Session = Depends(get_db)):
    """读取需求列宽配置"""
    proj = resolve_project(db, project_id)
    path = _col_widths_path(proj)
    if not os.path.isfile(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return {}


@router.put("/column-widths")
def save_column_widths(project_id: str, data: dict, db: Session = Depends(get_db)):
    """保存需求列宽配置"""
    proj = resolve_project(db, project_id)
    path = _col_widths_path(proj)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return {"ok": True}


@router.delete("/column-widths")
def delete_column_widths(project_id: str, db: Session = Depends(get_db)):
    """删除需求列宽配置（导入后重置）"""
    proj = resolve_project(db, project_id)
    path = _col_widths_path(proj)
    if os.path.isfile(path):
        os.remove(path)
    return {"ok": True}


# ========== 视图状态（排序/筛选）持久化 ==========

VIEW_STATE_FILE = "requirement_view.json"


def _view_state_path(proj):
    d = os.path.join(CONFIG_DIR, proj.display_id)
    return os.path.join(d, VIEW_STATE_FILE)


@router.get("/view-state")
def get_view_state(project_id: str, db: Session = Depends(get_db)):
    """读取排序和筛选配置"""
    proj = resolve_project(db, project_id)
    path = _view_state_path(proj)
    if not os.path.isfile(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return {}


@router.put("/view-state")
def save_view_state(project_id: str, data: dict, db: Session = Depends(get_db)):
    """保存排序和筛选配置"""
    proj = resolve_project(db, project_id)
    path = _view_state_path(proj)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return {"ok": True}


# ---- 富文本图片上传 ----

@router.post("/{requirement_id}/images")
def upload_requirement_image(
    project_id: str,
    requirement_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """上传富文本编辑器中的图片，返回可访问的 URL"""
    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)

    # 按项目显示ID/需求显示ID 分目录存储
    req_display_id = req.display_id or f"req_{req.id}"
    img_dir = os.path.join(UPLOAD_DIR, proj.display_id, "requirements", req_display_id, "images")
    os.makedirs(img_dir, exist_ok=True)

    # 生成唯一文件名
    ext = os.path.splitext(file.filename)[1] if file.filename else ".png"
    filename = f"{uuid.uuid4().hex}{ext}"
    filepath = os.path.join(img_dir, filename)

    content = file.file.read()
    with open(filepath, "wb") as f:
        f.write(content)

    # 返回可访问的 URL（相对 /uploads）
    url = f"/uploads/{proj.display_id}/requirements/{req_display_id}/images/{filename}"
    return {"url": url, "errno": 0}


@router.delete("/{requirement_id}/images/{filename}")
def delete_requirement_image(
    project_id: str,
    requirement_id: str,
    filename: str,
    db: Session = Depends(get_db),
):
    """删除富文本编辑器中的图片文件"""
    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)

    req_display_id = req.display_id or f"req_{req.id}"
    img_dir = os.path.join(UPLOAD_DIR, proj.display_id, "requirements", req_display_id, "images")
    filepath = os.path.join(img_dir, filename)

    if os.path.exists(filepath):
        os.remove(filepath)

    return {"ok": True}


# ---- 需求文件上传（超链接插入文件） ----

@router.post("/{requirement_id}/files")
def upload_requirement_file(
    project_id: str,
    requirement_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """上传文件作为超链接，返回可访问的 URL"""
    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)

    req_display_id = req.display_id or f"req_{req.id}"
    file_dir = os.path.join(UPLOAD_DIR, proj.display_id, "requirements", req_display_id, "files")
    os.makedirs(file_dir, exist_ok=True)

    # 保留原始扩展名
    ext = os.path.splitext(file.filename)[1] if file.filename else ""
    unique_name = f"{uuid.uuid4().hex}{ext}"
    filepath = os.path.join(file_dir, unique_name)

    size = 0
    max_size = 50 * 1024 * 1024  # 50MB
    content = file.file.read(1024 * 64)
    with open(filepath, "wb") as f:
        while content:
            size += len(content)
            if size > max_size:
                f.close()
                os.remove(filepath)
                raise HTTPException(413, "文件超过 50MB 限制")
            f.write(content)
            content = file.file.read(1024 * 64)

    url = f"/uploads/{proj.display_id}/requirements/{req_display_id}/files/{unique_name}"
    return {
        "url": url,
        "filename": unique_name,
        "original_filename": file.filename,
    }


@router.delete("/{requirement_id}/files/{filename}")
def delete_requirement_file(
    project_id: str,
    requirement_id: str,
    filename: str,
    db: Session = Depends(get_db),
):
    """删除需求超链接关联的文件"""
    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)

    req_display_id = req.display_id or f"req_{req.id}"
    file_dir = os.path.join(UPLOAD_DIR, proj.display_id, "requirements", req_display_id, "files")
    filepath = os.path.join(file_dir, filename)

    if os.path.exists(filepath):
        os.remove(filepath)

    return {"ok": True}


@router.get("/{requirement_id}/files/{filename}/preview")
def preview_requirement_file(
    project_id: str,
    requirement_id: str,
    filename: str,
    db: Session = Depends(get_db),
):
    """预览需求附件：Office 转 PDF、文本渲染为 HTML，其他 inline 展示"""
    from fastapi.responses import FileResponse as FR, HTMLResponse
    from urllib.parse import quote
    import html as html_mod
    from ..office_convert import is_office_file, convert_to_pdf

    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)

    req_display_id = req.display_id or f"req_{req.id}"
    file_dir = os.path.join(UPLOAD_DIR, proj.display_id, "requirements", req_display_id, "files")
    file_path = os.path.join(file_dir, filename)

    if not os.path.isfile(file_path):
        raise HTTPException(404, "文件不存在")

    # Office 文档 → 转换为 PDF 后预览
    if is_office_file(file_path):
        pdf_path = convert_to_pdf(file_path, file_dir)
        if pdf_path:
            return FR(
                pdf_path, media_type="application/pdf",
                headers={
                    "Content-Disposition": f"inline; filename*=UTF-8''{quote(filename, safe='')}",
                    "Cache-Control": "no-cache",
                }
            )
        else:
            return FR(
                file_path,
                media_type="application/octet-stream",
                headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename, safe='')}"}
            )

    ext = os.path.splitext(file_path)[1].lower()
    text_exts = {'.txt', '.log', '.md', '.sql', '.py', '.js', '.ts', '.html', '.css',
                 '.json', '.xml', '.yaml', '.yml', '.ini', '.cfg', '.conf',
                 '.sh', '.bat', '.ps1', '.csv', '.env', '.gitignore', '.dockerfile',
                 '.vue', '.java', '.c', '.cpp', '.h', '.go', '.rs', '.rb', '.php'}
    if ext in text_exts:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        html_content = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="utf-8"><title>{html_mod.escape(filename)}</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{background:#f8f8f8;padding:16px;font-family:'Cascadia Code','Consolas',monospace;font-size:14px;line-height:1.7}}
pre{{background:#fff;border:1px solid #e0e0e0;border-radius:6px;padding:16px;white-space:pre-wrap;word-break:break-all;color:#333}}
</style></head>
<body><pre>{html_mod.escape(content)}</pre></body></html>'''
        return HTMLResponse(
            content=html_content,
            headers={"Content-Disposition": f"inline; filename*=UTF-8''{quote(filename, safe='')}"}
        )

    return FR(
        file_path,
        headers={
            "Content-Disposition": f"inline; filename*=UTF-8''{quote(filename, safe='')}",
            "Cache-Control": "no-cache",
        }
    )


@router.get("/{requirement_id}", response_model=RequirementOut)
def get_requirement(
    project_id: str,
    requirement_id: str,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)
    # 重新加载以 eager load custom_values
    req = db.query(Requirement).options(
        joinedload(Requirement.custom_values).joinedload(RequirementCustomValue.field)
    ).filter(Requirement.id == req.id).first()
    return _format_requirement(req)


@router.put("/{requirement_id}", response_model=RequirementOut)
def update_requirement(
    project_id: str,
    requirement_id: str,
    data: RequirementUpdate,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)

    if data.title is not None:
        req.title = data.title
    if data.description is not None:
        req.description = data.description
    if data.priority is not None:
        req.priority = data.priority
    if data.status is not None:
        req.status = data.status

    # 更新自定义字段值
    _custom_values_changed = False
    if data.custom_values is not None:
        for field_id_str, value in data.custom_values.items():
            try:
                fid = int(field_id_str)
            except ValueError:
                continue
            existing = db.query(RequirementCustomValue).filter(
                RequirementCustomValue.requirement_id == req.id,
                RequirementCustomValue.field_id == fid,
            ).first()
            if existing:
                if value is not None and value != "":
                    existing.value = str(value)
                    _custom_values_changed = True
                else:
                    db.delete(existing)
                    _custom_values_changed = True
            else:
                if value is not None and value != "":
                    cv = RequirementCustomValue(
                        requirement_id=req.id,
                        field_id=fid,
                        value=str(value),
                    )
                    db.add(cv)
                    _custom_values_changed = True

    # 自定义字段变动时手动刷新 updated_at（ORM onupdate 无法感知关联表变动）
    if _custom_values_changed:
        req.updated_at = datetime.now()

    db.commit()
    db.refresh(req)
    touch_project(db, proj.id)
    return _get_requirement_with_values(db, req.id)


@router.delete("/{requirement_id}")
def delete_requirement(
    project_id: str,
    requirement_id: str,
    db: Session = Depends(get_db),
):
    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)
    db.delete(req)
    db.commit()
    touch_project(db, proj.id)
    return {"message": "ok"}


# ── 导出 ────────────────────────────────────────────────

def generate_requirement_doc_bytes(req, proj, db) -> bytes:
    """
    生成需求的 DOCX 文档字节（公文风格），供导出端点和任务导出共用。
    依赖 req.custom_values 已 eager load。
    """
    from ..database import RequirementStatusPool, RequirementPriorityPool, UPLOAD_DIR
    from ..export_service import (
        _add_run, _new_paragraph, _set_run_font, _set_heading_style,
        _setup_numbering, _apply_numbering, _set_cell_shading,
        _apply_table_widths, _add_h1,
        FONT_FAMILY, FONT_FAMILY_HEADING,
        BODY_SIZE, SMALL_SIZE, HEADING1_SIZE, HEADING2_SIZE, TITLE_SIZE, SUBTITLE_SIZE,
    )
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 获取池颜色
    status_pools = {p.name: p.color for p in db.query(RequirementStatusPool).filter(
        RequirementStatusPool.project_id == proj.id, RequirementStatusPool.is_active == True
    ).all()}
    priority_pools = {p.name: p.color for p in db.query(RequirementPriorityPool).filter(
        RequirementPriorityPool.project_id == proj.id, RequirementPriorityPool.is_active == True
    ).all()}

    # 状态/优先级 英文→中文映射
    STATUS_LABEL_MAP = {
        'todo': '待处理', 'in_progress': '进行中', 'done': '已完成', 'cancelled': '已取消',
    }
    PRIORITY_LABEL_MAP = {
        'low': '低', 'normal': '普通', 'high': '高', 'urgent': '紧急',
    }

    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = FONT_FAMILY
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY)
    pPr = style.element.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:line'), '360')
    pSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpacing)

    # 设置标题样式
    _set_heading_style(doc, 1, FONT_FAMILY_HEADING, HEADING1_SIZE)
    _set_heading_style(doc, 2, FONT_FAMILY_HEADING, HEADING2_SIZE)

    # 建立自动编号
    num_id = _setup_numbering(doc)

    # ---- 封面 ----
    for _ in range(6):
        _new_paragraph(doc, '', size=BODY_SIZE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _new_paragraph(doc, '需求说明文档', size=TITLE_SIZE, bold=True,
                   font_name=FONT_FAMILY_HEADING, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   before=200, after=100)
    _new_paragraph(doc, req.title or '(无标题)', size=SUBTITLE_SIZE,
                   font_name=FONT_FAMILY_HEADING, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   before=100, after=200)
    _new_paragraph(doc, '', size=BODY_SIZE)
    _new_paragraph(doc, f'项目名称：{proj.name}（{proj.display_id}）', size=SMALL_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _new_paragraph(doc, f'导出日期：{datetime.now().strftime("%Y年%m月%d日")}', size=SMALL_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    meta_items = [
        ("需求名称", req.title or "-"),
        ("显示ID", req.display_id or "-"),
        ("状态", STATUS_LABEL_MAP.get(req.status, req.status or "-")),
        ("优先级", PRIORITY_LABEL_MAP.get(req.priority, req.priority or "-")),
        ("创建时间", req.created_at.strftime("%Y-%m-%d %H:%M") if req.created_at else "-"),
        ("更新时间", req.updated_at.strftime("%Y-%m-%d %H:%M") if req.updated_at else "-"),
    ]

    # 收集自定义字段
    custom_items = []
    if req.custom_values:
        for cv in req.custom_values:
            if cv.field and cv.value:
                custom_items.append((cv.field.field_name, cv.value))

    _add_h1(doc, '需求基本信息', num_id, 0)
    _apply_table_widths(doc, meta_items)

    if custom_items:
        doc.add_paragraph()
        _add_h1(doc, '自定义字段', num_id, 0)
        _apply_table_widths(doc, custom_items)

    # ---- 详细描述 ----
    if req.description:
        _add_h1(doc, '详细描述', num_id, 0)
        _render_html_to_docx(doc, req.description, status_pools, priority_pools, img_base_dir=UPLOAD_DIR)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/{requirement_id}/export")
def export_requirement_doc(
    project_id: str,
    requirement_id: str,
    db: Session = Depends(get_db),
):
    """导出需求信息为 ZIP 压缩包（DOCX + 附件文件）"""
    import zipfile
    from ..database import Requirement, RequirementCustomField, RequirementCustomValue

    proj = resolve_project(db, project_id)
    req = resolve_requirement(db, proj.id, requirement_id)
    # 重新加载以 eager load custom_values
    req = db.query(Requirement).options(
        joinedload(Requirement.custom_values).joinedload(RequirementCustomValue.field)
    ).filter(Requirement.id == req.id).first()

    doc_bytes = generate_requirement_doc_bytes(req, proj, db)

    # 创建 ZIP 压缩包
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        # 写入 DOCX
        doc_name = f'{req.title}_需求文档_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
        zf.writestr(doc_name, doc_bytes)

        # 收集描述中引用的附件文件
        if req.description:
            file_pattern = re.compile(
                r'/uploads/' + re.escape(proj.display_id) +
                r'/requirements/[^/]+/files/([^"\s)]+)'
            )
            seen = set()
            for m in file_pattern.finditer(req.description):
                fn = m.group(1)
                if fn in seen:
                    continue
                seen.add(fn)
                file_path = os.path.join(
                    UPLOAD_DIR, proj.display_id, 'requirements',
                    req.display_id or f'req_{req.id}', 'files', fn
                )
                if os.path.isfile(file_path):
                    zf.write(file_path, f'files/{fn}')

    zip_bytes = zip_buf.getvalue()

    file_ts = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f'{req.title}_{file_ts}.zip'
    encoded_filename = urllib.parse.quote(filename)

    return Response(
        content=zip_bytes,
        media_type='application/zip',
        headers={
            'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}",
            'Content-Length': str(len(zip_bytes)),
        }
    )


def _css_bg_to_shd_fill(color_str: str) -> str:
    """
    将 CSS 背景色解析为 RGB，与白色混合变浅后返回六位十六进制颜色码。
    用于 run 级 w:shd 底纹填充（不限于 DOCX 16 种预设高亮色）。
    """
    if not color_str:
        return None
    c = color_str.strip()
    if c.startswith('#'):
        c = c.lstrip('#')
        if len(c) == 3:
            c = ''.join(x*2 for x in c)
        try:
            r, g, b = int(c[0:2],16), int(c[2:4],16), int(c[4:6],16)
        except:
            return None
    else:
        m = re.match(r'rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', c)
        if m:
            r, g, b = int(m.group(1)), int(m.group(2)), int(m.group(3))
        else:
            return None
    # 与白色混合使其变浅（factor 越大越浅）
    factor = 0.55
    r = int(r + (255 - r) * factor)
    g = int(g + (255 - g) * factor)
    b = int(b + (255 - b) * factor)
    return f'{r:02X}{g:02X}{b:02X}'


# ── HTML 描述转 DOCX ────────────────────────────────────

def _parse_inline_style(style_str: str) -> dict:
    """解析 style="key1:val1;key2:val2" 为字典"""
    result = {}
    if not style_str:
        return result
    for part in style_str.split(';'):
        part = part.strip()
        if ':' in part:
            k, v = part.split(':', 1)
            result[k.strip()] = v.strip()
    return result

def _color_to_rgb(color_str: str):
    """将 CSS 颜色转为 RGBColor，支持 #hex 和 rgb()"""
    from docx.shared import RGBColor
    color_str = color_str.strip()
    if color_str.startswith('#'):
        c = color_str.lstrip('#')
        if len(c) == 3:
            c = ''.join(x*2 for x in c)
        try:
            return RGBColor(int(c[0:2],16), int(c[2:4],16), int(c[4:6],16))
        except:
            return None
    m = re.search(r'rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_str)
    if m:
        return RGBColor(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    return None

def _darken_color(hex_color: str, factor: float = 0.6) -> str:
    """将十六进制颜色按 factor 加深（factor < 1 则变暗）。
    返回例如 '#b3b3b3' 的值，用于从背景色推导左侧边框色。"""
    c = hex_color.lstrip('#')
    if len(c) != 6:
        return '#888888'
    try:
        r = int(int(c[0:2], 16) * factor)
        g = int(int(c[2:4], 16) * factor)
        b = int(int(c[4:6], 16) * factor)
        return f'#{r:02x}{g:02x}{b:02x}'
    except ValueError:
        return '#888888'

def _cn(n):
    """阿拉伯数字转中文数字（1-99），用于标题一级编号 一、二、三"""
    if n <= 0:
        return '0'
    digits = '零一二三四五六七八九'
    if n < 10:
        return digits[n]
    if n < 20:
        return '十' + (digits[n % 10] if n % 10 else '')
    if n < 100:
        t, o = divmod(n, 10)
        return digits[t] + '十' + (digits[o] if o else '')
    return str(n)


def _render_html_to_docx(doc, html: str, status_pools: dict = None, priority_pools: dict = None, img_base_dir: str = None):
    """
    将 WangEditor 生成的 HTML 描述渲染到 docx 文档中，
    尽可能复刻 Web 上看到的效果。
    支持：加粗/斜体/下划线/删除线、字体颜色/背景色、引用块、列表、
         代码块、分割线、超链接、换行、图片。
    """
    from html.parser import HTMLParser
    from docx.shared import RGBColor, Inches, Pt as PtSize
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from ..export_service import FONT_FAMILY, FONT_FAMILY_HEADING, BODY_SIZE, SMALL_SIZE, HEADING1_SIZE, HEADING2_SIZE, _set_run_font, _add_hyperlink, _apply_numbering

    # 建立列表编号定义（bullet + decimal）
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part.element
    # 使用独立 numId: 70=bullet, 71=decimal（避开99占用）
    _NUM_BULLET = 70
    _NUM_DECIMAL = 71

    def _setup_list_num(num_id, fmt, text_pattern, start=1, levels=4):
        """为列表建立抽象编号定义，返回 num_id"""
        for n in numbering_elem.findall(qn('w:num')):
            if n.get(qn('w:numId')) == str(num_id):
                return num_id
        ab_id = str(1000 + num_id)  # 抽象编号 ID 不与现有冲突
        ab = OxmlElement('w:abstractNum')
        ab.set(qn('w:abstractNumId'), ab_id)
        for ilvl in range(levels):
            lvl = OxmlElement('w:lvl')
            lvl.set(qn('w:ilvl'), str(ilvl))
            for tag, val in [
                ('start', str(start)),
                ('numFmt', fmt),
                ('lvlText', text_pattern),
            ]:
                el = OxmlElement(f'w:{tag}')
                el.set(qn('w:val'), val)
                lvl.append(el)
            if fmt == 'bullet':
                # 子弹头符号设置
                numFont = OxmlElement('w:lvlJc')
                numFont.set(qn('w:val'), 'left')
                lvl.append(numFont)
                # 使用标准子弹头
                pStyle = OxmlElement('w:pStyle')
                lvl.append(pStyle)
            # 缩进
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(480 + ilvl * 480))
            ind.set(qn('w:hanging'), '240')
            lvl.append(ind)
            ab.append(lvl)
        numbering_elem.append(ab)
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_id))
        ref = OxmlElement('w:abstractNumId')
        ref.set(qn('w:val'), ab_id)
        num.append(ref)
        numbering_elem.append(num)
        return num_id

    _setup_list_num(_NUM_BULLET, 'bullet', '\u2022', levels=4)    # 实心圆点
    _setup_list_num(_NUM_DECIMAL, 'decimal', '%1.', levels=4)     # 1. 2. 3.

    class _HtmlRenderer(HTMLParser):
        def __init__(self, doc, img_base_dir=None):
            super().__init__()
            self.doc = doc
            self.img_base_dir = img_base_dir
            self.stack = []           # 标签栈
            self._p_texts = []        # 当前段落内 (text, bold, italic, underline, strikethrough, color, bg_color, link_url, font_name)
            self._in_bq = False
            self._bq_bg = None
            self._bq_border = None  # 引用块左侧边框色
            self._bq_table = None   # 引用块表格
            self._bq_cell = None    # 引用块单元格
            self._list_type = None
            self._list_depth = 0
            self._list_num_id = None
            self._skip_p = False
            self._code_mode = False    # 在 <pre>/<code> 内部
            self._heading_level = None   # 当前标题级别（1-6），None 表示非标题
            self._h_counters = [0, 0, 0]  # 多级标题计数器 h1/h2/h3
            self._in_caption = False      # 当前文本是否处于图注（span.req-caption）内
            # 表格支持
            self._in_td = False           # 当前是否在 <td>/<th> 内
            self._cell_runs = []          # 当前单元格段落的 run 列表
            self._cell_paragraphs = []    # 当前单元格的段落列表（每个段落 = run 列表）
            self._cell_is_header = False  # 当前单元格是否为表头
            self._cur_row = None          # 当前行（单元格列表）
            self._table = None            # 顶层表格缓冲 [row, ...]
            self._table_depth = 0         # 表格嵌套深度（仅渲染顶层）

        def _push_run(self, text='', bold=False, italic=False, underline=False, strikethrough=False,
                      color=None, bg_color=None, link_url='', font_name=None):
            if text:
                if self._in_td:
                    self._cell_runs.append((text, bold, italic, underline, strikethrough, color, bg_color, link_url, font_name))
                else:
                    self._p_texts.append((text, bold, italic, underline, strikethrough, color, bg_color, link_url, font_name))

        def _flush_paragraph(self):
            if self._in_td:
                return
            if not self._p_texts and not self._skip_p:
                return
            if self._skip_p:
                self._skip_p = False
                return

            # 图注：居中、灰色、斜体小字
            if self._in_caption:
                self._in_caption = False
                p = self.doc.add_paragraph()
                p.alignment = 1
                for text, bold, italic, underline, strikethrough, color, bg_color, link_url, font_name in self._p_texts:
                    run = p.add_run(text)
                    _set_run_font(run, size=SMALL_SIZE, color=_color_to_rgb('#888888'))
                    run.italic = True
                self._p_texts = []
                return


            p = (self._bq_cell.add_paragraph() if self._in_bq and self._bq_cell
                 else self.doc.add_paragraph())
            pPr = p._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            pPr.append(spacing)

            # 列表编号
            if self._list_type and self._list_num_id is not None:
                _apply_numbering(p, self._list_num_id, max(0, self._list_depth - 1))

            for text, bold, italic, underline, strikethrough, color, bg_color, link_url, font_name in self._p_texts:
                fn = font_name or FONT_FAMILY
                if link_url:
                    _add_hyperlink(p, text, link_url, size=BODY_SIZE)
                else:
                    run = p.add_run(text)
                    kwargs = {'bold': bold, 'size': BODY_SIZE, 'font_name': fn}
                    if italic:
                        run.italic = True
                    if underline:
                        run.underline = True
                    if strikethrough:
                        run.font.strike = True
                    if color:
                        rgb = _color_to_rgb(color)
                        if rgb:
                            kwargs['color'] = rgb
                    if bg_color:
                        fill_hex = _css_bg_to_shd_fill(bg_color)
                        if fill_hex:
                            rPr = run._r.get_or_add_rPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:fill'), fill_hex)
                            shd.set(qn('w:val'), 'clear')
                            rPr.append(shd)
                    _set_run_font(run, **kwargs)
            self._p_texts = []

        def _heading_label(self, level):
            """根据当前多级计数器生成标题编号文本：一、/ 1.1 / 1.1.1"""
            c = self._h_counters
            if level == 1:
                return _cn(c[0]) + '、'
            if level == 2:
                return f'{c[0]}.{c[1]} '
            return f'{c[0]}.{c[1]}.{c[2]} '

        def _flush_heading(self):
            """将当前累积文本作为多级编号标题输出（黑体、加粗）"""
            if not self._p_texts:
                self._p_texts = []
                return
            level = self._heading_level or 3
            # 更新多级计数器（进入更高级别时重置更深层）
            if level == 1:
                self._h_counters[0] += 1
                self._h_counters[1] = 0
                self._h_counters[2] = 0
            elif level == 2:
                self._h_counters[1] += 1
                self._h_counters[2] = 0
            elif level == 3:
                self._h_counters[2] += 1

            if level == 1:
                size = HEADING1_SIZE
            elif level == 2:
                size = HEADING2_SIZE
            else:
                size = PtSize(13)
            fn = FONT_FAMILY_HEADING

            p = self.doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '200')
            spacing.set(qn('w:after'), '120')
            spacing.set(qn('w:line'), '300')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)

            # 仅 1-3 级带自动编号前缀
            if 1 <= level <= 3:
                label = self._heading_label(level)
                run = p.add_run(label)
                _set_run_font(run, size=size, bold=True, font_name=fn)

            for text, bold, italic, underline, strikethrough, color, bg_color, link_url, font_name in self._p_texts:
                if link_url:
                    _add_hyperlink(p, text, link_url, size=size)
                else:
                    run = p.add_run(text)
                    _set_run_font(run, size=size, bold=True, font_name=fn)
            self._p_texts = []

        def _flush_cell_paragraph(self):
            """将当前单元格累积的 run 提交为一个段落（单元格内多段用 <p> 分隔）"""
            if self._cell_runs:
                self._cell_paragraphs.append(self._cell_runs)
                self._cell_runs = []

        def _flush_table(self):
            """将缓冲的表格数据渲染为 python-docx 表格"""
            rows = self._table or []
            if not rows:
                return
            n_cols = max((len(r) for r in rows), default=0)
            if n_cols == 0:
                return
            n_rows = len(rows)
            table = self.doc.add_table(rows=n_rows, cols=n_cols)
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
            table.autofit = True
            for i, row in enumerate(rows):
                for j in range(n_cols):
                    cell_paragraphs, is_header = row[j] if j < len(row) else ([], False)
                    cell = table.cell(i, j)
                    first = True
                    for para_runs in (cell_paragraphs or []):
                        p = cell.paragraphs[0] if first else cell.add_paragraph()
                        first = False
                        for (text, bold, italic, underline, strikethrough, color, bg_color, link_url, font_name) in para_runs:
                            if not text:
                                continue
                            run = p.add_run(text)
                            if bold or is_header:
                                run.bold = True
                            if italic:
                                run.italic = True
                            if underline:
                                run.underline = True
                            if strikethrough:
                                run.font.strike = True
                            if color:
                                rgb = _color_to_rgb(color)
                                if rgb:
                                    run.font.color.rgb = rgb
                            _set_run_font(run, size=BODY_SIZE, font_name=font_name or FONT_FAMILY)

        def _get_style_color(self, attrs_dict):
            """从 style 中提取 color 和 background-color"""
            style = attrs_dict.get('style', '')
            style_map = _parse_inline_style(style)
            return style_map.get('color', ''), style_map.get('background-color', '')

        def _add_image(self, attrs_dict):
            """将 <img> 标签嵌入 DOCX"""
            src = attrs_dict.get('src', '')
            if not src:
                return
            # 解析图片路径：/uploads/{project}/requirements/{req}/images/{filename}
            path = None
            if src.startswith('/uploads/'):
                if self.img_base_dir:
                    rel = src[len('/uploads/'):].replace('/', os.sep)
                    candidate = os.path.join(self.img_base_dir, rel)
                    if os.path.isfile(candidate):
                        path = candidate
            elif src.startswith('http://') or src.startswith('https://'):
                return
            if not path or not os.path.isfile(path):
                return
            try:
                self._flush_paragraph()
                p = (self._bq_cell.add_paragraph() if self._in_bq and self._bq_cell
                     else self.doc.add_paragraph())
                p.alignment = 1  # 居中
                run = p.add_run()
                style_map = _parse_inline_style(attrs_dict.get('style', ''))
                w_str = style_map.get('width', '')
                max_w = Inches(5.5)
                if w_str:
                    try:
                        w_px = float(w_str.replace('px', '').strip())
                        max_w = Inches(min(w_px / 96, 5.5))
                    except ValueError:
                        pass
                run.add_picture(path, width=max_w)
                self._skip_p = True
            except Exception:
                pass

        def handle_starttag(self, tag, attrs):
            attrs_dict = dict(attrs)
            if tag in ('p', 'div'):
                if self._in_td:
                    self._flush_cell_paragraph()
                else:
                    self._flush_paragraph()
                self.stack.append(tag)
            elif tag == 'table':
                self._flush_paragraph()
                self._table_depth += 1
                if self._table_depth == 1:
                    self._table = []
                    self._cur_row = None
                self.stack.append(tag)
            elif tag == 'tbody':
                self.stack.append(tag)
            elif tag == 'tr':
                self._cur_row = []
                self.stack.append(tag)
            elif tag in ('td', 'th'):
                self._in_td = True
                self._cell_runs = []
                self._cell_paragraphs = []
                self._cell_is_header = (tag == 'th')
                self.stack.append(tag)
            elif tag in ('h1','h2','h3','h4','h5','h6'):
                self._flush_paragraph()
                self._heading_level = int(tag[1])
                self.stack.append(tag)
            elif tag in ('b','strong','em','i','u','s','del','strike'):
                self.stack.append(tag)
            elif tag == 'blockquote':
                self._flush_paragraph()
                self._in_bq = True
                # 提取引用块颜色
                style = attrs_dict.get('style', '')
                style_map = _parse_inline_style(style)
                bg = style_map.get('background-color', '')
                if bg:
                    self._bq_bg = bg
                else:
                    bq_color = attrs_dict.get('data-bq-color', '')
                    if bq_color:
                        # bq_colors 可能为空（CSS 不在保存的 HTML 中）→ 直接用属性值
                        self._bq_bg = bq_colors.get(bq_color, '') or bq_color
                    else:
                        self._bq_bg = '#f0f0f0'

                # 提取边框颜色：data-bq-border > CSS 解析 > 背景色加深
                bq_border_raw = attrs_dict.get('data-bq-border', '')
                if bq_border_raw:
                    self._bq_border = bq_border_raw
                else:
                    # 尝试从 CSS 解析 border-left-color
                    bq_border_from_css = bq_border_colors.get(attrs_dict.get('data-bq-color', ''), '')
                    if bq_border_from_css:
                        self._bq_border = bq_border_from_css
                    else:
                        # 回退：背景色加深为边框色
                        self._bq_border = _darken_color(self._bq_bg) if self._bq_bg else '#888888'

                # 创建单格表格模拟引用块
                bq_color_hex = self._bq_bg.replace('#', '') if self._bq_bg else 'f0f0f0'
                bq_border_hex = self._bq_border.replace('#', '') if self._bq_border else bq_color_hex

                table = self.doc.add_table(rows=1, cols=1)
                table.autofit = True

                # 取单元格
                cell = table.rows[0].cells[0]
                self._bq_cell = cell

                # 删除默认空段落（否则第一行前会多一个回车）
                default_p = cell.paragraphs[0]._element
                default_p.getparent().remove(default_p)

                # 设置单元格宽度占满 + 清除表格级边框
                tblPr = table._tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._tbl.insert(0, tblPr)
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '5000')
                tblW.set(qn('w:type'), 'pct')
                tblPr.append(tblW)
                # 清除表格级边框
                tblBorders = OxmlElement('w:tblBorders')
                for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    sideEl = OxmlElement(f'w:{side}')
                    sideEl.set(qn('w:val'), 'none')
                    sideEl.set(qn('w:sz'), '0')
                    sideEl.set(qn('w:space'), '0')
                    sideEl.set(qn('w:color'), 'auto')
                    tblBorders.append(sideEl)
                tblPr.append(tblBorders)

                # 单元格底纹
                if self._bq_bg:
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), bq_color_hex)
                    shd.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shd)

                # 设置单元格边框：左侧 12pt 加粗彩色竖线，上右下无边框
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for side, sz, color, val in [
                    ('top',    '0',  'auto',            'none'),   # 无边框
                    ('left',   '24', bq_border_hex,     'single'), # 12pt 彩色竖线
                    ('bottom', '0',  'auto',            'none'),   # 无边框
                    ('right',  '0',  'auto',            'none'),   # 无边框
                ]:
                    sideEl = OxmlElement(f'w:{side}')
                    sideEl.set(qn('w:val'), val)
                    sideEl.set(qn('w:sz'), sz)
                    sideEl.set(qn('w:space'), '4')
                    sideEl.set(qn('w:color'), color)
                    tcBorders.append(sideEl)
                tcPr.append(tcBorders)

                # 边距（5pt = 100dxa 上下留白）
                tcMar = OxmlElement('w:tcMar')
                for side, val in [('top', '100'), ('left', '120'), ('bottom', '100'), ('right', '60')]:
                    mar = OxmlElement(f'w:{side}')
                    mar.set(qn('w:w'), val)
                    mar.set(qn('w:type'), 'dxa')
                    tcMar.append(mar)
                tcPr.append(tcMar)

                self.stack.append(tag)
            elif tag == 'hr':
                self._flush_paragraph()
                # 分割线：底部边框的段落
                p = self.doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '4')
                bottom.set(qn('w:color'), '999999')
                pBdr.append(bottom)
                pPr.append(pBdr)
            elif tag in ('ul', 'ol'):
                if tag == 'ul':
                    self._list_type = 'ul'
                    self._list_num_id = _NUM_BULLET
                else:
                    self._list_type = 'ol'
                    self._list_num_id = _NUM_DECIMAL
                self._list_depth += 1
                self.stack.append(tag)
            elif tag == 'li':
                self._flush_paragraph()
                self.stack.append(tag)
            elif tag == 'a':
                self.stack.append(('a', attrs_dict.get('href', '')))
            elif tag == 'br':
                self._push_run('\n')
            elif tag == 'img':
                if self._in_td:
                    return
                self._add_image(attrs_dict)
            elif tag == 'span':
                self.stack.append(('span', attrs_dict))
            elif tag == 'pre':
                self._flush_paragraph()
                self._code_mode = True
                self.stack.append(tag)
            elif tag == 'code':
                self._code_mode = True
                self.stack.append(('code', attrs_dict.get('style', '')))

        def handle_endtag(self, tag):
            if tag in ('p', 'div'):
                if self._in_td:
                    self._flush_cell_paragraph()
                else:
                    self._flush_paragraph()
                self._pop_stack(tag)
            elif tag == 'table':
                self._table_depth -= 1
                if self._table_depth == 0:
                    self._flush_table()
                    self._table = None
                self._pop_stack(tag)
            elif tag == 'tbody':
                self._pop_stack(tag)
            elif tag == 'tr':
                if self._cur_row:
                    self._table.append(self._cur_row)
                self._cur_row = None
                self._pop_stack(tag)
            elif tag in ('td', 'th'):
                self._flush_cell_paragraph()
                self._cur_row.append((self._cell_paragraphs, self._cell_is_header))
                self._in_td = False
                self._cell_runs = []
                self._cell_paragraphs = []
                self._cell_is_header = False
                self._pop_stack(tag)
            elif tag in ('h1','h2','h3','h4','h5','h6'):
                self._flush_heading()
                self._heading_level = None
                self._pop_stack(tag)
            elif tag in ('b','strong','em','i','u','s','del','strike'):
                self._pop_stack(tag)
            elif tag == 'blockquote':
                self._flush_paragraph()
                self._in_bq = False
                self._bq_bg = None
                self._bq_border = None
                self._bq_cell = None
                self._pop_stack(tag)
            elif tag in ('ul', 'ol'):
                if tag == 'ul':
                    self._list_type = 'ul' if self._has_type_above('ul') else None
                else:
                    self._list_type = 'ol' if self._has_type_above('ol') else None
                self._list_depth = max(0, self._list_depth - 1)
                self._pop_stack(tag)
            elif tag == 'li':
                self._flush_paragraph()
                self._skip_p = True
                self._pop_stack(tag)
            elif isinstance(self.stack[-1] if self.stack else None, tuple) and self.stack[-1][0] == 'a' and tag == 'a':
                self._pop_stack(tag)
            elif tag == 'span':
                self._pop_stack(tag)
            elif tag == 'pre':
                self._code_mode = False
                self._flush_paragraph()
                self._pop_stack(tag)
            elif tag == 'code':
                self._code_mode = False
                self._pop_stack(tag)

        def _pop_stack(self, tag):
            if self.stack and self.stack[-1] == tag:
                self.stack.pop()
            elif self.stack and isinstance(self.stack[-1], tuple) and self.stack[-1][0] == tag:
                self.stack.pop()

        def _has_type_above(self, target):
            for item in reversed(self.stack[:-1]):
                if item == target:
                    return True
                if item == ('ul' if target == 'ol' else 'ol'):
                    return False
            return False

        def handle_data(self, data):
            if not data.strip() and not self._code_mode:
                return

            bold = any(t in self.stack for t in ('b', 'strong'))
            italic = any(t in self.stack for t in ('i', 'em'))
            underline = 'u' in self.stack
            strikethrough = any(t in self.stack for t in ('s', 'del', 'strike'))
            link_url = ''
            color = ''
            bg_color = ''
            font_name = None

            for item in self.stack:
                if isinstance(item, tuple):
                    if item[0] == 'a':
                        link_url = item[1]
                        # 需求附件文件 → 转为 ZIP 内相对路径
                        m = re.match(r'^/uploads/[^/]+/requirements/[^/]+/files/(.+)$', link_url)
                        if m:
                            link_url = 'files/' + m.group(1)
                        # 补全缺少协议的链接（如 www.baidu.com → https://www.baidu.com）
                        elif link_url and not re.match(r'^[a-zA-Z][a-zA-Z0-9+\-.]*://', link_url) and not link_url.startswith('/'):
                            link_url = 'https://' + link_url
                    elif item[0] == 'span':
                        attrs = item[1]
                        c, bg = self._get_style_color(attrs)
                        if c: color = c
                        if bg: bg_color = bg
                        cls = attrs.get('class', '') or ''
                        if 'req-caption' in cls and not self._in_td:
                            self._in_caption = True
                            color = '#888888'
                    elif item[0] == 'code':
                        font_name = 'Courier New'
                        bg_color = '#cccccc'

            if self._code_mode:
                font_name = 'Courier New'
                bg_color = '#cccccc'

            self._push_run(data, bold=bold, italic=italic, underline=underline,
                          strikethrough=strikethrough, color=color, bg_color=bg_color,
                          link_url=link_url, font_name=font_name)

    # 清理 + 提取全局 CSS 中的 blockquote 颜色
    bq_colors = {}
    bq_border_colors = {}
    for m in re.finditer(r'blockquote\[data-bq-color=["\']([^"\']+)["\']\]\s*\{[^}]*background-color:\s*([^;}]+)', html):
        bq_colors[m.group(1)] = m.group(2).strip()
    for m in re.finditer(r'blockquote\[data-bq-color=["\']([^"\']+)["\']\][^{]*\{[^}]*background[^:]*:\s*([^;}]+)', html):
        bq_colors[m.group(1)] = m.group(2).strip()
    # 同时提取 border-left-color
    for m in re.finditer(r'blockquote\[data-bq-color=["\']([^"\']+)["\']\][^{]*\{[^}]*border-left-color:\s*([^;}]+)', html):
        bq_border_colors[m.group(1)] = m.group(2).strip()

    # 兜底：所有未着色的 blockquote 自动继承前一个有色 blockquote 的颜色
    # （处理 WangEditor 多段引用块未正确注入 data-bq-color 的情况）
    html = _propagate_bq_colors(html)

    html = html.strip()
    if html.startswith('<html>') or html.startswith('<!DOCTYPE'):
        m = re.search(r'<body[^>]*>([\s\S]*)</body>', html, re.I)
        if m:
            html = m.group(1)

    renderer = _HtmlRenderer(doc, img_base_dir=img_base_dir)
    renderer.feed(html)
    renderer._flush_paragraph()


def _propagate_bq_colors(html: str) -> str:
    """对 HTML 中未着色的 <blockquote> 向前/后传播继承最近的有色块。
    纯字符串处理，对抗 WangEditor 序列化丢失属性的情况。"""
    if 'blockquote' not in html:
        return html
    # 提取所有 <blockquote ...> 标签及属性
    bq_pattern = re.compile(r'<blockquote(\s[^>]*)?>', re.IGNORECASE)
    matches = list(bq_pattern.finditer(html))
    if len(matches) < 2:
        return html
    # 提取每段的 color/border 属性
    def get_attrs(tag_text: str):
        c = re.search(r'data-bq-color=["\']([^"\']+)["\']', tag_text or '')
        b = re.search(r'data-bq-border=["\']([^"\']+)["\']', tag_text or '')
        return (c.group(1) if c else None, b.group(1) if b else None)
    # 先按"原位属性"建立数组
    bq_attrs = [get_attrs(m.group(1)) for m in matches]
    # 双向传播
    last_color = None
    last_border = None
    for i in range(len(bq_attrs)):
        c, b = bq_attrs[i]
        if c:
            last_color, last_border = c, b
        elif last_color:
            bq_attrs[i] = (last_color, last_border)
    last_color = None
    last_border = None
    for i in range(len(bq_attrs) - 1, -1, -1):
        c, b = bq_attrs[i]
        if c:
            last_color, last_border = c, b
        elif last_color:
            bq_attrs[i] = (last_color, last_border)
    # 重新拼接（仅修改有变化的）
    result = []
    last_end = 0
    for m, (c, b) in zip(matches, bq_attrs):
        if not c:
            result.append(html[last_end:m.end()])
            last_end = m.end()
            continue
        original_attrs = m.group(1) or ''
        # 在原 attrs 基础上补充/覆盖 color/border
        if 'data-bq-color=' in original_attrs:
            new_attrs = re.sub(r'data-bq-color=["\'][^"\']*["\']', f'data-bq-color="{c}"', original_attrs)
        else:
            new_attrs = original_attrs.rstrip() + f' data-bq-color="{c}"'
        if b:
            if 'data-bq-border=' in new_attrs:
                new_attrs = re.sub(r'data-bq-border=["\'][^"\']*["\']', f'data-bq-border="{b}"', new_attrs)
            else:
                new_attrs = new_attrs.rstrip() + f' data-bq-border="{b}"'
        new_tag = f'<blockquote{new_attrs}>'
        result.append(html[last_end:m.start()])
        result.append(new_tag)
        last_end = m.end()
    result.append(html[last_end:])
    return ''.join(result)


# ── Excel 导入 ─────────────────────────────────────────
import openpyxl
from io import BytesIO
from pydantic import BaseModel
from typing import List, Dict, Optional
from fastapi import UploadFile, File, Form

# ──────── 通用 Excel 解析 ────────
def _cell_str(v) -> str:
    """将 Excel 单元格值安全转为字符串，修复 lone surrogate"""
    if v is None:
        return ""
    s = str(v).strip()
    try:
        s.encode("utf-8")
    except UnicodeEncodeError:
        s = s.encode("utf-8", errors="surrogatepass").decode("utf-8", errors="replace").strip()
    return s


def _parse_excel(content: bytes, filename: str = "") -> tuple:
    """
    解析 Excel 内容，返回 (headers, all_rows)。
    headers — 第一行各列文本
    all_rows — 从第二行开始的所有数据行（每行与 headers 等长，空串填充）
    """
    is_xls = filename.lower().endswith(".xls") and not filename.lower().endswith(".xlsx")

    if is_xls:
        # .xls 优先用 xlrd，失败时 fallback 到 openpyxl
        try:
            return _parse_xls(content)
        except HTTPException:
            raise
        except Exception as e_xls:
            try:
                return _parse_xlsx(content)
            except HTTPException:
                raise
            except Exception:
                raise HTTPException(400, f"无法解析 .xls 文件: {e_xls}")
    else:
        # .xlsx 优先用 openpyxl，失败时 fallback 到 xlrd
        try:
            return _parse_xlsx(content)
        except HTTPException:
            raise
        except Exception as e_xlsx:
            try:
                return _parse_xls(content)
            except HTTPException:
                raise
            except Exception:
                raise HTTPException(400, f"无法解析 .xlsx 文件: {e_xlsx}")


def _parse_xlsx(content: bytes) -> tuple:
    """用 openpyxl 解析 .xlsx 内容"""
    wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
    ws = wb.active
    if ws is None:
        wb.close()
        raise HTTPException(400, "Excel 文件没有工作表")

    # 扫描所有有数据的单元格，自行探测行列范围
    # 不依赖 ws.max_column/max_row，因为部分 WPS/非标准 Excel 文件的
    # dimension 可能返回错误值，导致 iter_rows 丢列或空迭代。
    cell_data = {}  # (row, col) -> value
    max_row, max_col = 0, 0
    for (r, c), cell in ws._cells.items():
        if cell.value is not None:
            cell_data[(r, c)] = cell.value
            max_row = max(max_row, r)
            max_col = max(max_col, c)

    if max_row == 0:
        wb.close()
        raise HTTPException(400, "Excel 文件没有数据")

    if max_col == 0:
        wb.close()
        raise HTTPException(400, "Excel 文件没有检测到列")

    # 第一行作表头
    headers = [_cell_str(cell_data.get((1, col))) for col in range(1, max_col + 1)]

    all_rows = []
    for r in range(2, max_row + 1):
        row_vals = [_cell_str(cell_data.get((r, col))) for col in range(1, max_col + 1)]
        all_rows.append(row_vals)

    wb.close()
    return headers, all_rows


def _parse_xls(content: bytes) -> tuple:
    """用 xlrd 解析 .xls 内容（旧版 Excel 格式）"""
    import xlrd
    wb = xlrd.open_workbook(file_contents=content)
    ws = wb.sheet_by_index(0)
    if ws.nrows == 0:
        raise HTTPException(400, "Excel 文件为空")

    first = [str(ws.cell_value(0, c)).strip() for c in range(ws.ncols)]
    all_rows = []
    for r in range(1, ws.nrows):
        row_vals = [str(ws.cell_value(r, c)).strip() if ws.cell_type(r, c) != xlrd.XL_CELL_EMPTY else "" for c in range(ws.ncols)]
        all_rows.append(row_vals)
    return first, all_rows


class ExcelPreviewOut(BaseModel):
    headers: List[str]
    rows: List[List[str]]
    total_rows: int


@router.post("/import/preview")
async def import_preview(
    project_id: str,
    file: UploadFile = File(...),
):
    """上传 Excel 文件，预览表头和前几行数据"""
    content = await file.read()
    if not content:
        raise HTTPException(400, "上传的文件为空")

    filename = file.filename or ""
    try:
        headers, all_rows = _parse_excel(content, filename)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"无法解析 Excel 文件: {e}")

    preview_rows = all_rows[:5]
    return ExcelPreviewOut(
        headers=headers,
        rows=preview_rows,
        total_rows=len(all_rows),
    )


class ColumnMapping(BaseModel):
    target: str  # "title" | "status" | "priority" | "field:{id}" | "new"
    field_name: Optional[str] = None
    field_type: Optional[str] = None
    field_options: Optional[str] = None


def _find_title_col(excel_headers, mapping_dict):
    """从 mapping_dict 中找到标题列的索引"""
    for h, m_obj in mapping_dict.items():
        if m_obj.target == "title" and h in excel_headers:
            return excel_headers.index(h)
    return None


def _detect_file_duplicates(all_rows, title_col):
    """检测文件内重复标题，返回 [(title, [row_numbers])]"""
    title_groups = defaultdict(list)
    for idx, row_vals in enumerate(all_rows):
        if not any(row_vals):
            continue
        val = (row_vals[title_col] or "").strip() if title_col < len(row_vals) else ""
        if val:
            title_groups[val].append(idx + 2)  # +2 因为 Excel 行号（表头占1）
    return [(t, r) for t, r in title_groups.items() if len(r) > 1]


def _dedup_rows_add_sequence(all_rows, title_col):
    """添加序号策略：重复标题加 _1 _2 后缀，全部保留"""
    counter = defaultdict(int)
    result = []
    for row_vals in all_rows:
        title = (row_vals[title_col] or "").strip() if title_col < len(row_vals) else ""
        if title:
            counter[title] += 1
            if counter[title] > 1:
                new_title = f"{title}_{counter[title] - 1}"
                row_vals[title_col] = new_title
        result.append(row_vals)
    return result


@router.post("/import")
async def import_requirements(
    project_id: str,
    file: UploadFile = File(...),
    mapping: str = Form(...),  # JSON: {"Excel列名": ColumnMapping}
    mode: str = Form("append"),  # append | overwrite | update
    force: bool = Form(False),  # True = 跳过重复检测
    dup_strategy: str = Form("cancel"),  # cancel | add_sequence
    db: Session = Depends(get_db),
):
    """根据列映射导入 Excel 数据到需求
    mode: append=直接新增, overwrite=清空后重新导入, update=标题匹配更新/新增
    """
    import json
    mapping_dict: Dict[str, ColumnMapping] = {}
    try:
        raw = json.loads(mapping)
        for k, v in raw.items():
            mapping_dict[k] = ColumnMapping(**v)
    except Exception as e:
        raise HTTPException(400, f"映射格式错误: {e}")

    if not mapping_dict:
        raise HTTPException(400, "未配置任何列映射")

    proj = resolve_project(db, project_id)
    content = await file.read()
    if not content:
        raise HTTPException(400, "上传的文件为空")

    filename = file.filename or ""
    try:
        excel_headers, all_rows = _parse_excel(content, filename)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"无法解析 Excel 文件: {e}")

    # ── 重复检测与冲突处理 ──
    if not force:
        title_col = _find_title_col(excel_headers, mapping_dict)
        dup_result = None  # {message, dialog_type, actions, file_duplicates}

        if title_col is not None:
            file_dups = _detect_file_duplicates(all_rows, title_col)

            if mode == 'overwrite' and file_dups:
                # overwrite: 文件内重复 → 二选一
                dup_result = {
                    "file_duplicates": [{"title": t, "rows": r} for t, r in file_dups],
                    "dialog_type": "choice",
                    "actions": ["cancel", "add_sequence"],
                    "message": f"文件内发现 {len(file_dups)} 组重复标题",
                }

            elif mode == 'update' and file_dups:
                # update: 检查重复标题是否在 DB 中有对应
                dup_titles = [t for t, _ in file_dups]
                db_existing = set(
                    r[0] for r in db.query(Requirement.title).filter(
                        Requirement.project_id == proj.id,
                        Requirement.title.in_(dup_titles),
                    ).all()
                )
                db_conflict_dups = [(t, r) for t, r in file_dups if t in db_existing]
                no_conflict_dups = [(t, r) for t, r in file_dups if t not in db_existing]

                if db_conflict_dups:
                    # 有冲突的重复 → 不可处理
                    if no_conflict_dups:
                        dup_result = {
                            "file_duplicates": [{"title": t, "rows": r} for t, r in file_dups],
                            "dialog_type": "abandon_only",
                            "actions": ["cancel"],
                            "message": (
                                f"以下重复标题已存在于数据库，无法匹配更新："
                                + "、".join(t for t, _ in db_conflict_dups)
                                + "。请取消导入并调整文件。"
                            ),
                        }
                    else:
                        dup_result = {
                            "file_duplicates": [{"title": t, "rows": r} for t, r in file_dups],
                            "dialog_type": "abandon_only",
                            "actions": ["cancel"],
                            "message": "重复标题已存在于数据库，无法匹配更新，请取消导入并调整文件。",
                        }
                elif no_conflict_dups:
                    # 纯重复，DB 中无对应 → 二选一
                    dup_result = {
                        "file_duplicates": [{"title": t, "rows": r} for t, r in no_conflict_dups],
                        "dialog_type": "choice",
                        "actions": ["cancel", "add_sequence"],
                        "message": f"文件内发现 {len(no_conflict_dups)} 组重复标题（数据库中无对应）",
                    }

            elif mode == 'append' and file_dups:
                # append: 检测冲突+非冲突
                dup_titles = [t for t, _ in file_dups]
                db_existing = set(
                    r[0] for r in db.query(Requirement.title).filter(
                        Requirement.project_id == proj.id,
                        Requirement.title.in_(dup_titles),
                    ).all()
                )
                db_conflict_dups = [(t, r) for t, r in file_dups if t in db_existing]
                no_conflict_dups = [(t, r) for t, r in file_dups if t not in db_existing]

                if db_conflict_dups and not no_conflict_dups:
                    # 全冲突 → 提示信息
                    dup_result = {
                        "file_duplicates": [{"title": t, "rows": r} for t, r in db_conflict_dups],
                        "dialog_type": "info_only",
                        "actions": ["ok"],
                        "message": f"重复标题「{'、'.join(t for t, _ in db_conflict_dups)}」已存在于数据库，已跳过，无需额外处理。",
                    }
                elif no_conflict_dups:
                    parts = []
                    if db_conflict_dups:
                        parts.append(f"「{'、'.join(t for t, _ in db_conflict_dups)}」→ 已存在于DB，已跳过")
                    # 纯重复部分需要用户选择
                    dup_result = {
                        "file_duplicates": [{"title": t, "rows": r} for t, r in no_conflict_dups],
                        "dialog_type": "choice",
                        "actions": ["cancel", "add_sequence"],
                        "message": "文件内发现重复标题" + ("，" + "；".join(parts) if parts else ""),
                    }

        if dup_result:
            dup_result["warning"] = True
            return dup_result

    # ── 按策略去重 ──
    if dup_strategy == "add_sequence":
        title_col = _find_title_col(excel_headers, mapping_dict)
        if title_col is not None:
            all_rows = _dedup_rows_add_sequence(all_rows, title_col)

    # ── 覆盖模式：先清空所有已有数据，再重新创建 ──
    if mode == 'overwrite':
        db.query(RequirementCustomValue).filter(
            RequirementCustomValue.requirement_id.in_(
                db.query(Requirement.id).filter(Requirement.project_id == proj.id)
            )
        ).delete(synchronize_session=False)
        db.query(Requirement).filter(Requirement.project_id == proj.id).delete(synchronize_session=False)
        # 删除非内置自定义字段（含旧版导入遗留的与内置同名的字段）
        db.query(RequirementCustomField).filter(
            RequirementCustomField.project_id == proj.id,
            RequirementCustomField.is_builtin == False,
        ).delete(synchronize_session=False)
        db.query(RequirementStatusPool).filter(
            RequirementStatusPool.project_id == proj.id
        ).delete(synchronize_session=False)
        db.query(RequirementPriorityPool).filter(
            RequirementPriorityPool.project_id == proj.id
        ).delete(synchronize_session=False)
        db.flush()

    # 解析映射并创建新字段
    field_id_cache: Dict[str, Optional[int]] = {}

    def resolve_field_target(excel_col: str) -> Optional[int]:
        if excel_col not in mapping_dict:
            return None
        m = mapping_dict[excel_col]
        if m.target.startswith("field:"):
            try:
                return int(m.target.split(":")[1])
            except (ValueError, IndexError):
                return None
        if m.target != "new":
            return None
        # 新字段名称若匹配内置字段名，改为映射到内置列，不创建新字段
        field_name = m.field_name or excel_col
        builtin_name_map = {"标题": "title", "状态": "status", "优先级": "priority"}
        if field_name in builtin_name_map:
            mapping_dict[excel_col].target = builtin_name_map[field_name]
            return None
        cf = RequirementCustomField(
            project_id=proj.id,
            field_name=field_name,
            field_type=m.field_type or "text",
            field_options=m.field_options or "",
            sort_order=999,
        )
        db.add(cf)
        db.flush()
        return cf.id

    for h in excel_headers:
        if h in mapping_dict:
            field_id_cache[h] = resolve_field_target(h)
        else:
            field_id_cache[h] = None

    # 枚举值规范化：中文/英文混合输入都能识别
    status_name_to_value = {"待处理": "todo", "进行中": "in_progress", "已完成": "done", "已取消": "cancelled"}
    status_value_set = {"todo", "in_progress", "done", "cancelled"}
    priority_name_to_value = {"低": "low", "普通": "normal", "高": "high", "紧急": "urgent"}
    priority_value_set = {"low", "normal", "high", "urgent"}

    def _normalize_status(v: str) -> str:
        v = (v or "").strip()
        if not v:
            return "todo"
        if v in status_value_set:
            return v
        if v in status_name_to_value:
            return status_name_to_value[v]
        # 检查是否在状态池中（含刚同步添加的）
        in_pool = db.query(RequirementStatusPool).filter(
            RequirementStatusPool.project_id == proj.id,
            RequirementStatusPool.name == v,
            RequirementStatusPool.is_active == True,
        ).first()
        if in_pool:
            return v
        return "todo"

    def _normalize_priority(v: str) -> str:
        v = (v or "").strip()
        if not v:
            return "normal"
        if v in priority_value_set:
            return v
        if v in priority_name_to_value:
            return priority_name_to_value[v]
        # 检查是否在优先级池中（含刚同步添加的）
        in_pool = db.query(RequirementPriorityPool).filter(
            RequirementPriorityPool.project_id == proj.id,
            RequirementPriorityPool.name == v,
            RequirementPriorityPool.is_active == True,
        ).first()
        if in_pool:
            return v
        return "normal"

    # ── 同步池数据：将导入值自动添加到状态池/优先级池/下拉选项 ──
    def _sync_import_pools():

        # 查找 Excel 中哪一列映射到了 status
        status_col = None
        priority_col = None
        drop_field_ids = set()
        for h, m in mapping_dict.items():
            if m.target == 'status':
                status_col = h
            elif m.target == 'priority':
                priority_col = h
            elif m.target.startswith('field:'):
                try:
                    fid = int(m.target.split(':')[1])
                except (ValueError, IndexError):
                    continue
                if fid in drop_field_ids:
                    continue
                cf = db.query(RequirementCustomField).filter(
                    RequirementCustomField.id == fid,
                    RequirementCustomField.project_id == proj.id,
                ).first()
                if cf and cf.field_type in ('dropdown', 'multi_dropdown'):
                    drop_field_ids.add(fid)
        # 此外，field_id_cache 中新建（new）的字段也可能是 dropdown
        for h, fid in field_id_cache.items():
            if fid is None:
                continue
            if fid in drop_field_ids:
                continue
            if h not in mapping_dict:
                continue
            if mapping_dict[h].target != 'new':
                continue
            cf = db.query(RequirementCustomField).filter(
                RequirementCustomField.id == fid,
                RequirementCustomField.project_id == proj.id,
            ).first()
            if cf and cf.field_type in ('dropdown', 'multi_dropdown'):
                drop_field_ids.add(fid)

        # 收集唯一值
        status_vals = set()
        priority_vals = set()
        drop_vals = {}  # field_id -> set of values
        for row_vals in all_rows:
            for i, h in enumerate(excel_headers):
                val = row_vals[i].strip() if i < len(row_vals) else ''
                if not val:
                    continue
                if h == status_col:
                    status_vals.add(val)
                elif h == priority_col:
                    priority_vals.add(val)
                elif h in mapping_dict:
                    mt = mapping_dict[h].target
                    fid = None
                    if mt.startswith('field:'):
                        try:
                            fid = int(mt.split(':')[1])
                        except (ValueError, IndexError):
                            continue
                    elif mt == 'new':
                        # 新建字段：从 field_id_cache 取 id
                        fid = field_id_cache.get(h)
                    if fid is not None and fid in drop_field_ids:
                        if fid not in drop_vals:
                            drop_vals[fid] = set()
                        for part in val.split(','):
                            p = part.strip()
                            if p:
                                drop_vals[fid].add(p)

        # 同步状态池
        existing_statuses = {
            r.name for r in db.query(RequirementStatusPool).filter(
                RequirementStatusPool.project_id == proj.id
            ).all()
        }
        for v in sorted(status_vals):
            if v not in existing_statuses and v not in status_value_set:
                sp = RequirementStatusPool(
                    project_id=proj.id, name=v, color='#5F5E5A', sort_order=999
                )
                db.add(sp)

        # 同步优先级池
        existing_priorities = {
            r.name for r in db.query(RequirementPriorityPool).filter(
                RequirementPriorityPool.project_id == proj.id
            ).all()
        }
        for v in sorted(priority_vals):
            if v not in existing_priorities and v not in priority_value_set:
                pp = RequirementPriorityPool(
                    project_id=proj.id, name=v, color='#5F5E5A', sort_order=999
                )
                db.add(pp)

        # 同步下拉/多选自定义字段的选项
        for fid, val_set in drop_vals.items():
            cf = db.query(RequirementCustomField).filter(
                RequirementCustomField.id == fid,
                RequirementCustomField.project_id == proj.id,
            ).first()
            if not cf:
                continue
            try:
                opts = json.loads(cf.field_options) if cf.field_options else []
            except (json.JSONDecodeError, TypeError):
                opts = []
            existing_labels = {o.get('label', '') for o in opts}
            added = []
            for v in sorted(val_set):
                if v not in existing_labels:
                    opts.append({'label': v, 'color': '#5F5E5A'})
                    added.append(v)
            if added:
                cf.field_options = json.dumps(opts, ensure_ascii=False)
                db.add(cf)

    _sync_import_pools()
    db.flush()

    # ── 根据模式处理已有数据 ──
    updated = 0
    existing_titles = set()  # append 模式用于检测 DB 冲突
    if mode == 'update':
        # 更新模式：预加载所有已有需求的标题→id映射
        existing = {
            r.title: r for r in db.query(Requirement).filter(
                Requirement.project_id == proj.id
            ).all()
        }
    elif mode == 'append':
        # 追加模式：只加载标题用于冲突检测
        existing_titles = set(
            r[0] for r in db.query(Requirement.title).filter(
                Requirement.project_id == proj.id
            ).all()
        )

    created = 0
    skipped_empty_title = 0
    skipped_db_collision = 0
    for row_vals in all_rows:
        if not any(row_vals):
            continue

        data = {}
        custom_vals = {}

        for i, h in enumerate(excel_headers):
            val = row_vals[i] if i < len(row_vals) else ""
            if h not in mapping_dict:
                continue
            m = mapping_dict[h]

            if m.target in ("title", "status", "priority"):
                data[m.target] = val
            elif field_id_cache[h] is not None:
                custom_vals[str(field_id_cache[h])] = val

        # 标题必填：空标题整行跳过
        raw_title = (data.get("title") or "").strip()
        if not raw_title:
            skipped_empty_title += 1
            continue

        # 追加模式：冲突标题跳过
        if mode == 'append' and raw_title in existing_titles:
            skipped_db_collision += 1
            continue

        if mode == 'update' and raw_title in existing:
            # 更新已有需求
            req = existing[raw_title]
            changed = False
            if data.get('priority'):
                req.priority = _normalize_priority(data.get('priority'))
                changed = True
            if data.get('status'):
                req.status = _normalize_status(data.get('status'))
                changed = True

            # 更新自定义字段值
            for fid_str, val in custom_vals.items():
                if not val:
                    continue
                existing_cv = db.query(RequirementCustomValue).filter(
                    RequirementCustomValue.requirement_id == req.id,
                    RequirementCustomValue.field_id == int(fid_str),
                ).first()
                if existing_cv:
                    if existing_cv.value != val:
                        existing_cv.value = val
                        db.add(existing_cv)
                        changed = True
                else:
                    db.add(RequirementCustomValue(
                        requirement_id=req.id,
                        field_id=int(fid_str),
                        value=val,
                    ))
                    changed = True

            if changed:
                req.updated_at = datetime.now()
                db.add(req)
                db.flush()
            updated += 1
        else:
            # 新增需求
            req = Requirement(
                project_id=proj.id,
                display_id=generate_requirement_display_id(db, proj),
                title=raw_title,
                priority=_normalize_priority(data.get('priority')),
                status=_normalize_status(data.get('status')),
            )
            db.add(req)
            db.flush()

            for fid_str, val in custom_vals.items():
                if val:
                    cv = RequirementCustomValue(
                        requirement_id=req.id,
                        field_id=int(fid_str),
                        value=val,
                    )
                    db.add(cv)
            created += 1

    _sync_builtin_field_options(proj, db)
    db.commit()
    touch_project(db, proj.id)
    msg_parts = []
    if created:
        msg_parts.append(f"新增 {created} 条")
    if updated:
        msg_parts.append(f"更新 {updated} 条")
    msg = "成功导入，" + "，".join(msg_parts) if msg_parts else "无数据变更"
    if skipped_empty_title:
        msg += f"，跳过 {skipped_empty_title} 条无标题空行"
    if skipped_db_collision:
        msg += f"，跳过 {skipped_db_collision} 条（数据库中已存在）"
    return {"created": created, "updated": updated, "skipped_empty": skipped_empty_title, "skipped_db_collision": skipped_db_collision, "message": msg}
